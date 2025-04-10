import os
from pathlib import Path
from typing import List, Union
from fpdf import FPDF
from PIL import Image
from docx import Document
from PyPDF2 import PdfReader, PdfWriter

def checkFilePermissions(inputPath: Path, outputPath: Path, needWrite: bool = False) -> None:
    """Check file permissions and existence."""
    if not inputPath.exists():
        raise FileNotFoundError(f"Input file not found: {inputPath}")
    if not os.access(inputPath, os.R_OK):
        raise PermissionError(f"No read permission for: {inputPath}")
    if needWrite and outputPath.exists() and not os.access(outputPath, os.W_OK):
        raise PermissionError(f"No write permission for existing file: {outputPath}")
    if not os.access(outputPath.parent, os.W_OK):
        raise PermissionError(f"No write permission for output directory: {outputPath.parent}")

def convertTextToPdf(textFilePath: Union[str, Path], outputPdfPath: Union[str, Path]) -> None:
    """
    Convert a text file to PDF.
    
    Args:
        textFilePath: Path to the input text file.
        outputPdfPath: Path where the output PDF will be saved.
        
    Raises:
        FileNotFoundError: If the input file doesn't exist.
        PermissionError: If there are permission issues.
        Exception: For other conversion errors.
    """
    try:
        textFile = Path(textFilePath)
        outputPdf = Path(outputPdfPath)
        
        checkFilePermissions(textFile, outputPdf, needWrite=True)
            
        with open(textFile, "r", encoding="utf-8") as txtFile:
            text = txtFile.read()
            
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.multi_cell(0, 10, text)
        pdf.output(outputPdf)
        print(f"Successfully converted text file to PDF: {outputPdf}")
        
    except (FileNotFoundError, PermissionError) as e:
        print(f"Access error: {e}")
        raise
    except Exception as e:
        print(f"Error converting text to PDF: {e}")
        raise

def convertPngToJpg(pngFilePath: Union[str, Path], outputJpgPath: Union[str, Path], quality: int = 85) -> None:
    """
    Convert a PNG image to JPG.
    
    Args:
        pngFilePath: Path to the input PNG file.
        outputJpgPath: Path where the output JPG will be saved.
        quality: Quality of the output JPG (1-100).
        
    Raises:
        FileNotFoundError: If the input file doesn't exist.
        PermissionError: If there are permission issues.
        ValueError: If quality is not between 1 and 100.
        Exception: For other conversion errors.
    """
    try:
        pngFile = Path(pngFilePath)
        outputJpg = Path(outputJpgPath)
        
        checkFilePermissions(pngFile, outputJpg, needWrite=True)
        if not 1 <= quality <= 100:
            raise ValueError("Quality must be between 1 and 100")
            
        with Image.open(pngFile) as img:
            if img.mode != 'RGB':
                img = img.convert("RGB")
            img.save(outputJpg, "JPEG", quality=quality)
        print(f"Successfully converted PNG to JPG: {outputJpg}")
        
    except (FileNotFoundError, PermissionError) as e:
        print(f"Access error: {e}")
        raise
    except Exception as e:
        print(f"Error converting PNG to JPG: {e}")
        raise

def convertWordToPdf(wordFilePath: Union[str, Path], outputPdfPath: Union[str, Path]) -> None:
    """
    Convert a Word document to PDF.
    
    Args:
        wordFilePath: Path to the input Word document.
        outputPdfPath: Path where the output PDF will be saved.
        
    Raises:
        FileNotFoundError: If the input file doesn't exist.
        PermissionError: If there are permission issues.
        Exception: For other conversion errors.
    """
    try:
        wordFile = Path(wordFilePath)
        outputPdf = Path(outputPdfPath)
        
        checkFilePermissions(wordFile, outputPdf, needWrite=True)
            
        doc = Document(wordFile)
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        
        for para in doc.paragraphs:
            pdf.multi_cell(0, 10, para.text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    pdf.multi_cell(0, 10, cell.text)
        
        pdf.output(outputPdf)
        print(f"Successfully converted Word file to PDF: {outputPdf}")
        
    except (FileNotFoundError, PermissionError) as e:
        print(f"Access error: {e}")
        raise
    except Exception as e:
        print(f"Error converting Word to PDF: {e}")
        raise

def mergePdfFiles(pdfFiles: List[Union[str, Path]], outputPdfPath: Union[str, Path]) -> None:
    """
    Merge multiple PDF files into one.
    
    Args:
        pdfFiles: List of paths to the input PDF files.
        outputPdfPath: Path where the output merged PDF will be saved.
        
    Raises:
        FileNotFoundError: If any input file doesn't exist.
        PermissionError: If there are permission issues.
        ValueError: If no PDF files are provided.
        Exception: For other merging errors.
    """
    try:
        if not pdfFiles:
            raise ValueError("No PDF files provided for merging")
            
        pdfPaths = [Path(pdf) for pdf in pdfFiles]
        outputPdf = Path(outputPdfPath)
        
        for pdfFile in pdfPaths:
            checkFilePermissions(pdfFile, outputPdf, needWrite=True)
                
        pdfWriter = PdfWriter()
        
        for pdfFile in pdfPaths:
            with open(pdfFile, "rb") as file:
                pdfReader = PdfReader(file)
                for page in pdfReader.pages:
                    pdfWriter.add_page(page)
        
        with open(outputPdf, "wb") as outputFile:
            pdfWriter.write(outputFile)
            
        print(f"Successfully merged {len(pdfPaths)} PDFs into: {outputPdf}")
        
    except (FileNotFoundError, PermissionError) as e:
        print(f"Access error: {e}")
        raise
    except Exception as e:
        print(f"Error merging PDFs: {e}")
        raise