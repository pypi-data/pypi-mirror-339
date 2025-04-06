import os
import re
import traceback
from pathlib import Path
from typing import List

from docx import Document
from lxml import etree as et
from natsort import natsorted

from criticus.py.reformat_collation.itsee_to_open_cbgm import reformat_xml

TEI_NS = "{http://www.tei-c.org/ns/1.0}"
XML_NS = "{http://www.w3.org/XML/1998/namespace}"
ABBR_TO_FULL = {
    "Matt": "Matthew",
    "B01": "Matthew",
    "Mark": "Mark",
    "B02": "Mark",
    "Luke": "Luke",
    "B03": "Luke",
    "John": "John",
    "B04": "John",
    "Acts": "Acts",
    "B05": "Acts",
    "Rom": "Romans",
    "B06": "Romans",
    "Romans": "Romans",
    "R": "Romans",
    "1 Cor": "1 Corinthians",
    "1Cor": "1 Corinthians",
    "ICor": "1 Corinthians",
    "B07": "1 Corinthians",
    "1 Corinthians": "1 Corinthians",
    "2 Cor": "2 Corinthians",
    "2Cor": "2 Corinthians",
    "IICor": "2 Corinthians",
    "2 Corinthians": "2 Corinthians",
    "B08": "2 Corinthians",
    "Gal": "Galatians",
    "Galatians": "Galatians",
    "B09": "Galatians",
    "Eph": "Ephesians",
    "Ephesians": "Ephesians",
    "B10": "Ephesians",
    "Phil": "Philippians",
    "Philippians": "Philippians",
    "B11": "Philippians",
    "Col": "Colossians",
    "Colossians": "Colossians",
    "B12": "Colossians",
    "1 Thess": "1 Thessalonians",
    "1Thess": "1 Thessalonians",
    "1 Thessalonians": "1 Thessalonians",
    "B13": "1 Thessalonians",
    "2 Thess": "2 Thessalonians",
    "2Thess": "2 Thessalonians",
    "2 Thessalonians": "2 Thessalonians",
    "B14": "2 Thessalonians",
    "1 Tim": "1 Timothy",
    "1Tim": "1 Timothy",
    "1 Timothy": "1 Timothy",
    "B15": "1 Timothy",
    "2 Tim": "2 Timothy",
    "2Tim": "2 Timothy",
    "2 Timothy": "2 Timothy",
    "B16": "2 Timothy",
    "Titus": "Titus",
    "B17": "Titus",
    "Phlm": "B18",
    "Philemon": "Philemon",
    "B18": "Philemon",
    "Heb": "B19",
    "Hebrews": "Hebrews",
    "B19": "Hebrews",
    "Jas": "James",
    "James": "James",
    "B20": "James",
    "1 Pet": "1 Peter",
    "1Pet": "1 Peter",
    "1 Peter": "1 Peter",
    "B21": "1 Peter",
    "2 Pet": "2 Peter",
    "2Pet": "2 Peter",
    "2 Peter": "2 Peter",
    "B22": "2 Peter",
    "1 John": "1 John",
    "B23": "1 John",
    "1John": "1 John",
    "2 John": "2 John",
    "B24": "2 John",
    "2John": "2 John",
    "3 John": "3John",
    "3John": "3John",
    "B25": "3John",
    "Jude": "Jude",
    "B26": "Jude",
    "Rev": "Revelation",
    "Revelation": "Revelation",
    "B27": "Revelation",
}


def get_xml_file(xml: str) -> tuple[et._Element, bool] | tuple[str, None]:
    temp_cx_file = "temp_xml_collation_file"
    # xml = xml.replace('xml:id="1', 'xml:id="I')
    # xml = xml.replace('xml:id="2', 'xml:id="II')
    # xml = xml.replace('xml:id="3', 'xml:id="III')
    xml = xml.replace("subreading", "subr")
    with open(temp_cx_file, "w", encoding="utf-8") as file:
        file.write(xml)
    if re.search("<teiHeader>", xml) is None:
        try:
            temp_cx_file = reformat_xml(temp_cx_file, "title", "publication")
        except Exception:
            return traceback.format_exc(), None
    parser = et.XMLParser(remove_blank_text=True, encoding="UTF-8", recover=True)
    tree = et.parse(temp_cx_file, parser)  # type: et._ElementTree
    root = tree.getroot()
    os.remove(temp_cx_file)
    return root, True


def get_document():
    this_dir = Path(__file__).parent
    template = this_dir.joinpath("template.docx").as_posix()
    # print(template)
    return Document(template)


def load_xml_file(xml_file: str):
    with open(xml_file, "r", encoding="utf-8") as file:
        xml = file.read()
    return get_xml_file(xml)


def construct_full_ref(ab: et._Element):
    ref = ab.get(f"{XML_NS}id").replace("-APP", "").upper()  # type: str
    if ref.startswith("B"):  # then it is an INTF/IGNTP style reference... probably
        book = re.search(r"B\d+", ref).group(0)
        book = ABBR_TO_FULL[book]
        chapter = re.search(r"K\d+", ref).group(0)
        verse = re.search(r"V\d+", ref).group(0)
        ref = f"{book} {chapter.replace('K', '')}:{verse.replace('V', '')}"
    else:
        book = re.search(r".[a-zA-Z]+", ref)
        if not book:
            return ref
        book = book.group(0)
        full_book = ABBR_TO_FULL.get(book)
        if not full_book:
            return ref
        reference = ref.replace(book, "").replace(".", ":")
        ref = f"{full_book} {reference}"
    return ref


def print_reference(document: Document, ab: et._Element):
    ref = construct_full_ref(ab)
    reference = document.add_paragraph(ref)
    reference.style = document.styles["reference"]


def group_basetext_words(basetext: str, words_per_line: int) -> List[list]:
    words_per_line = words_per_line - 1
    grouped_basetext = []
    current_group = []
    chunk = 0
    for word in basetext.split():
        if chunk == words_per_line:
            current_group.append(word)
            grouped_basetext.append(current_group)
            chunk = 0
            current_group = []
            continue
        current_group.append(word)
        chunk += 1
    if current_group != []:
        grouped_basetext.append(current_group)
    return grouped_basetext


def construct_basetext(ab: et._Element) -> str:
    if ab.text:
        return ab.text
    basetext = []
    for elem in ab:
        if elem.tag == f"{TEI_NS}seg":
            text = elem.text
            if text:
                basetext.append(text)
            else:
                print(f"seg element {elem.get(f'{TEI_NS}n')} has no text")
        elif (
            elem.tag == f"{TEI_NS}app" and elem.find(f"{TEI_NS}lem").get("type") != "om"
        ):
            text = elem.find(f"{TEI_NS}lem").text
            if text:
                basetext.append(text)
            else:
                print("lem text is None???")
    return " ".join(basetext)


def print_basetext(document: Document, ab: et._Element, words_per_line: int):
    basetext = construct_basetext(ab)
    basetext = group_basetext_words(basetext, words_per_line)
    table = document.add_table(rows=0, cols=10)
    index = 2
    for line in basetext:
        row_cells = table.add_row().cells
        for cell, word in enumerate(line):
            row_cells[cell].text = f"{word}\n{index}"
            row_cells[cell].paragraphs[0].style = document.styles["table cell"]
            index += 2


def print_app(document: Document, app: et._Element):
    app_from = app.get("from")
    app_to = app.get("to")
    if app_from == app_to:
        index = app_from
    else:
        index = f"{app_from}–{app_to}"
    p = document.add_paragraph(index)
    p.style = document.styles["index"]


def sort_by_ga(wits: List[str]):
    papyri = []
    majuscules = []
    minuscules = []
    lectionaries = []
    editions = []
    for wit in wits:
        # try:
        if wit.lower().startswith("p"):
            papyri.append(wit)
        elif wit.startswith("0"):
            majuscules.append(wit)
        elif wit[0].isdigit():
            minuscules.append(wit)
        elif wit.lower().startswith("l"):
            lectionaries.append(wit)
        else:
            editions.append(wit)
        # except:
        #     print(wit)
    return (
        natsorted(papyri)
        + natsorted(majuscules)
        + natsorted(minuscules)
        + natsorted(lectionaries)
        + natsorted(editions)
    )


# TODO: format wits same as Apparatus Explorer
# def format_wits(wits: str):
#     wits = re.sub(r'\([^()]\)', '', wits)
#     wits = wits.split()
#     wits = sort_by_ga(wits)


def print_rdg(
    document,
    rdg: et._Element,
    text_wits_separator: str,
    rdg_n_text_separator: str,
    text_bold: bool,
    wits_separator: str,
):
    if rdg.get("type") and rdg.text:
        rdg_type, greek_text = f"\t{rdg.get('type')}", rdg.text
    elif rdg.get("type"):
        rdg_type, greek_text = f"\t{rdg.get('type')}", ""
    else:
        rdg_type, greek_text = "\t", rdg.text
    # if rdg.text:
    #     greek_text = rdg.text
    # else:
    #     greek_text = rdg.get('type')
    p = document.add_paragraph()
    p.style = document.styles["reading"]
    rdg_name = re.sub(r"\d", "", rdg.get("n"))
    p.add_run(f"{rdg_name}.").italic = True
    p.add_run(f"{rdg_type}\t").italic = True
    p.add_run(rdg_n_text_separator)
    p.add_run(greek_text).bold = text_bold
    wits = rdg.get("wit")

    wits = sort_by_ga(wits.split())
    wits = wits_separator.join(wits)
    wits = wits.replace("(", " (").replace("_", " ")
    p.add_run(f"{text_wits_separator}{wits}")


def combine_regularized(app: et._Element, add_suffix: bool):
    reg_wits = {}
    for rdg in app.findall(f"{TEI_NS}rdg"):
        rdg_id = rdg.get("n")
        if "r" not in rdg_id:
            parent = rdg_id
        elif rdg_id[0] != parent[0]:
            continue
        else:
            reg_wits[parent] = f"{reg_wits.get(parent, '')} {rdg.get('wit')}".lstrip()

    for parent, wits in reg_wits.items():
        wits = wits.split()
        if add_suffix:
            wits = [f"{w}r" for w in wits]
        wits = " ".join(wits)
        reg_wits[parent] = wits

    previous = ""
    for rdg in app.findall(f"{TEI_NS}rdg"):
        if rdg.get("n") in reg_wits:
            combined = f"{rdg.get('wit')} {reg_wits[rdg.get('n')]}".lstrip()
            rdg.attrib["wit"] = combined
        elif "r" in rdg.get("n") and rdg.get("n")[0] != previous[0]:
            continue
        elif "r" in rdg.get("n"):
            app.remove(rdg)
        previous = rdg.get("n")

    return app


def export_xml_to_docx(
    xml_filename: str,
    output_filename: str,
    basetext_words_per_line: int,
    text_wits_separator: str,
    rdg_n_text_separator: str,
    text_bold: bool,
    wits_separator: str,
    custom_template: str,
    use_custom_template: bool = False,
    collapse_regularized: bool = False,
    add_suffix: bool = False,
):
    if use_custom_template:
        document = Document(custom_template)
    else:
        document = get_document()
    if not document:
        return False, "Could not load document template."
    result, success = load_xml_file(xml_filename)
    if not success:
        return False, result
    root = result
    for ab in root.findall(f"{TEI_NS}ab"):
        print_reference(document, ab)
        print_basetext(document, ab, basetext_words_per_line)
        for app in ab.findall(f"{TEI_NS}app"):
            if collapse_regularized:
                app = combine_regularized(app, add_suffix)
            if len(app.findall(f"{TEI_NS}rdg")) == 1:
                continue

            print_app(document, app)
            for rdg in app.findall(f"{TEI_NS}rdg"):
                print_rdg(
                    document,
                    rdg,
                    text_wits_separator,
                    rdg_n_text_separator,
                    text_bold,
                    wits_separator,
                )

    document.save(output_filename)
    return True, "Success"
