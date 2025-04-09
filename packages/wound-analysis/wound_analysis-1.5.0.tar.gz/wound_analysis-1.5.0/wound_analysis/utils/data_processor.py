import logging
import os
import pathlib
import traceback
from dataclasses import dataclass, field
from datetime import datetime
from typing import Any, Dict, Optional, Literal, List, TypeAlias
from pandas import Timestamp

import numpy as np
import math
from scipy import stats
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import streamlit as st

from wound_analysis.utils.column_schema import DColumns, ExcelSheetColumns


VisitsDataType: TypeAlias = Dict[ DColumns | Literal['wound_measurements', 'sensor_data'],
								Timestamp | Dict[ str, float | Dict[str, Any]]]

VisitsMetadataType: TypeAlias = Dict[Literal['patient_metadata', 'visits'], Dict | List[VisitsDataType]]



logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


# Helper function to convert NaN to None
def get_float(data, key):
	return float(data[key]) if (data is not None) and (not pd.isna(data.get(key))) else None


class WoundDataProcessor:
	def __init__(self, impedance_freq_sweep_path: pathlib.Path=None, df: pd.DataFrame=None, impedance_analyzer: Optional['ImpedanceAnalyzer']=None):
		"""Initialize the WoundDataProcessor.

		Args:
			impedance_freq_sweep_path: Path to impedance frequency sweep files
			df: Pre-loaded and preprocessed DataFrame
			impedance_analyzer: Optional existing ImpedanceAnalyzer instance to reuse
		"""
		self.impedance_freq_sweep_path = impedance_freq_sweep_path
		self.df = df.copy()

		# Use provided analyzer or create a new one
		if impedance_analyzer is not None:
			self.impedance_analyzer = impedance_analyzer
		else:
			# Initialize a new analyzer if none was provided
			self.impedance_analyzer = ImpedanceAnalyzer(impedance_freq_sweep_path=impedance_freq_sweep_path)

		# Initialize column schema
		self.CN = DColumns(df=self.df)

		self._cache_patient_visits = {}

		for record_id in self.df[self.CN.RECORD_ID].unique():
			_ = self.get_patient_visits(record_id=record_id)


	def get_cache_key_patient_visits(self, record_id: int) -> str:
		return f"{record_id}_visits"


	def get_patient_visits(self, record_id: int) -> VisitsMetadataType:
		"""
			Retrieves all visit data for a specific patient.

			This method fetches all visits for a patient identified by their record ID, processes
			the data for each non-skipped visit, and includes wound information.

			Args:
				record_id (int): The unique identifier for the patient.

			Returns:
				Dict: A dictionary containing:
					- 'patient_metadata': Basic patient information extracted from their first visit
					- 'visits': List of dictionaries containing data for each valid visit, including wound information

			Raises:
				ValueError: If no records are found for the specified patient
				Exception: If any error occurs during data processing

			Note:
				This method uses caching to improve performance when processing impedance data.
				The cache is used to avoid reprocessing the same data multiple times.
		"""

		cache_key = self.get_cache_key_patient_visits(record_id=record_id)

		if cache_key in self._cache_patient_visits:
			return self._cache_patient_visits[cache_key]

		# try:
		data_csv_for_patient = self.df[self.df[self.CN.RECORD_ID] == record_id]
		if data_csv_for_patient.empty:
			raise ValueError(f"No measurements found for patient {record_id}")

		# Get metadata from first visit
		first_visit = data_csv_for_patient.iloc[0]
		metadata = self._extract_patient_metadata(first_visit)

		visits_data = []

		for _, visit_csv_row in data_csv_for_patient.iterrows():

			if pd.isna(visit_csv_row.get(self.CN.SKIPPED_VISIT)) or visit_csv_row[self.CN.SKIPPED_VISIT] != 'Yes':

				# Check if we already have cached data for this visit
				visit_data = self._process_visit_data(visit_csv_row=visit_csv_row, record_id=record_id)

				if visit_data:

					wound_info = self._get_wound_info(visit_csv_row)

					visit_data['wound_info'] = wound_info

					visits_data.append(visit_data)

		result = { 'patient_metadata': metadata, 'visits': visits_data }

		self._cache_patient_visits[cache_key] = result

		return result


	def get_population_statistics(self) -> Dict:
		"""
			Gather comprehensive population-level statistics for LLM analysis.

			This method aggregates data from all patients and visits to provide a holistic view
			of the wound care dataset. It includes information from various aspects of wound care,
			patient demographics, and treatment outcomes.

			Returns:
				Dict: A comprehensive dictionary containing population statistics including:
					- Summary (total patients, visits, overall improvement rate, etc.)
					- Demographics (ag gender, race, BMI distribution)
					- Wound characteristics (types, locations, initial sizes)
					- Healing progression metrics (healing rates, time to closure)
					- Sensor data analysis (impedance trends, temperature patterns, oxygenation levels)
					- Risk factor analysis (comorbidities, listyle factors)
					- Treatment effectiveness (comparison of different approaches)
					- Temporal trends (seasonal variations, long-term outcome improvements)

			Note:
				This method processes the entire dataset and may be computationally intensive
				for large datasets.
		"""

		# Get processed data
		df = self.get_processed_data()
		if df.empty:
			raise ValueError("No data available for population statistics")

		CN = self.CN
		# Calculate derived metrics
		df[CN.BMI_CATEGORY] = pd.cut(
			df[CN.BMI],
			bins=[0, 18.5, 24.9, 29.9, float('inf')],
			labels=['Underweight', 'Normal', 'Overweight', 'Obese']
		)

		# Calculate healing metrics
		df['Healing_Color']  = df[CN.HEALING_RATE].apply(lambda x: 'green' if x < 0 else 'red')
		df['Healing_Status'] = df[CN.HEALING_RATE].apply(
			lambda x: 'Improving' if x < 0 else ('Stable' if -5 <= x <= 5 else 'Worsening')
		)

		stats_data = {
			'summary': {
				'total_patients'             : len(df[CN.RECORD_ID].unique()),
				'total_visits'               : len(df),
				'avg_visits_per_patient'     : len(df) / len(df[CN.RECORD_ID].unique()),
				'overall_improvement_rate'   : (df[CN.HEALING_RATE] < 0).mean() * 100,
				'avg_treatment_duration_days': (df.groupby(CN.RECORD_ID)[CN.DAYS_SINCE_FIRST_VISIT].max().mean()),
			},
			'demographics': {
				'age_stats': {
					'summary': f"Mean: {df[CN.AGE].mean():.1f}, Median: {df[CN.AGE].median():.1f}",
					'distribution': df[CN.AGE].value_counts().to_dict(),
					'age_groups': pd.cut(df[CN.AGE],
						bins=[0, 30, 50, 70, float('inf')],
						labels=['<30', '30-50', '50-70', '>70']).value_counts().to_dict()
				},
				'gender_distribution'   : df[CN.SEX].value_counts().to_dict(),
				'race_distribution'     : df[CN.RACE].value_counts().to_dict(),
				'ethnicity_distribution': df[CN.ETHNICITY].value_counts().to_dict(),
				'bmi_stats': {
					'summary': f"Mean: {df[CN.BMI].mean():.1f}, Range: {df[CN.BMI].min():.1f}-{df[CN.BMI].max():.1f}",
					'distribution'     : df[CN.BMI_CATEGORY].value_counts().to_dict(),
					'by_healing_status': df.groupby(CN.BMI_CATEGORY)[CN.HEALING_RATE].agg(['mean', 'count']).to_dict()
				}
			},
			'risk_factors': {
				'primary_conditions': {
					'diabetes': {
						'distribution'  : df[CN.DIABETES].value_counts().to_dict(),
						'healing_impact': df.groupby(CN.DIABETES)[CN.HEALING_RATE].agg(['mean', 'std', 'count']).to_dict()
					},
					'smoking': {
						'distribution'  : df[CN.SMOKING_STATUS].value_counts().to_dict(),
						'healing_impact': df.groupby(CN.SMOKING_STATUS)[CN.HEALING_RATE].agg(['mean', 'std', 'count']).to_dict()
					}
				},
				'comorbidity_analysis': {
					'diabetes_smoking': df.groupby([CN.DIABETES, CN.SMOKING_STATUS])[CN.HEALING_RATE].agg(['mean', 'count']).to_dict(),
					'diabetes_bmi'    : df.groupby([CN.DIABETES, CN.BMI_CATEGORY])[CN.HEALING_RATE].agg(['mean', 'count']).to_dict()
				}
			},
			'wound_characteristics': {
				'type_distribution': {
					'overall': df[CN.WOUND_TYPE].value_counts().to_dict(),
					'by_healing_status': df.groupby([CN.WOUND_TYPE, 'Healing_Status']).size().to_dict()
				},
				'location_analysis': {
					'distribution': df[CN.WOUND_LOCATION].value_counts().to_dict(),
					'healing_by_location': df.groupby(CN.WOUND_LOCATION)[CN.HEALING_RATE].mean().to_dict()
				},
				'size_progression': {
					'initial_vs_final': {
						'area': {
							'initial': df.groupby(CN.RECORD_ID)[CN.WOUND_AREA].first().agg(['mean', 'median', 'std']).to_dict(),
							'final'  : df.groupby(CN.RECORD_ID)[CN.WOUND_AREA].last().agg(['mean', 'median', 'std']).to_dict(),
							'percent_change': ((df.groupby(CN.RECORD_ID)[CN.WOUND_AREA].last() -
												df.groupby(CN.RECORD_ID)[CN.WOUND_AREA].first()) /
												df.groupby(CN.RECORD_ID)[CN.WOUND_AREA].first() * 100).mean()
						}
					},
					'healing_by_initial_size': {
						'small' : df[df[CN.WOUND_AREA] < df[CN.WOUND_AREA].quantile(0.33)][CN.HEALING_RATE].mean(),
						'medium': df[(df[CN.WOUND_AREA] >= df[CN.WOUND_AREA].quantile(0.33)) &
									(df[CN.WOUND_AREA] < df[CN.WOUND_AREA].quantile(0.67))][CN.HEALING_RATE].mean(),
						'large': df[df[CN.WOUND_AREA] >= df[CN.WOUND_AREA].quantile(0.67)][CN.HEALING_RATE].mean()
					}
				}
			},
			'healing_progression': {
				'overall_stats': {
					'summary': f"Mean: {df[CN.HEALING_RATE].mean():.1f}%, Median: {df[CN.HEALING_RATE].median():.1f}%",
					'distribution': df['Healing_Status'].value_counts().to_dict(),
					'percentiles' : df[CN.HEALING_RATE].quantile([0.25, 0.5, 0.75]).to_dict()
				},
				'temporal_analysis': {
					'by_visit_number': df.groupby(CN.VISIT_NUMBER)[CN.HEALING_RATE].agg(['mean', 'std', 'count']).to_dict(),
					'by_treatment_duration': pd.cut(df[CN.DAYS_SINCE_FIRST_VISIT],
						bins=[0, 30, 90, 180, float('inf')],
						labels=['<30 days', '30-90 days', '90-180 days', '>180 days']
					).value_counts().to_dict()
				}
			},
			'exudate_analysis': {
				'characteristics': {
					'volume': {
						'distribution'       : df[CN.EXUDATE_VOLUME].value_counts().to_dict(),
						'healing_correlation': df.groupby(CN.EXUDATE_VOLUME)[CN.HEALING_RATE].mean().to_dict()
					},
					'type': {
						'distribution'       : df[CN.EXUDATE_TYPE].value_counts().to_dict(),
						'healing_correlation': df.groupby(CN.EXUDATE_TYPE)[CN.HEALING_RATE].mean().to_dict()
					},
					'viscosity': {
						'distribution'       : df[CN.EXUDATE_VISCOSITY].value_counts().to_dict(),
						'healing_correlation': df.groupby(CN.EXUDATE_VISCOSITY)[CN.HEALING_RATE].mean().to_dict()
					}
				},
				'temporal_patterns': {
					'volume_progression': df.groupby(CN.VISIT_NUMBER)[CN.EXUDATE_VOLUME].value_counts().to_dict(),
					'type_progression'  : df.groupby(CN.VISIT_NUMBER)[CN.EXUDATE_TYPE].value_counts().to_dict()
				}
			}
		}

		# Add sensor data analysis if available
		stats_data['sensor_data'] = {}

		# Temperature Analysis
		if CN.CENTER_TEMP in df.columns:
			stats_data['sensor_data']['temperature'] = {
				'center_temp': {
					'overall': df[CN.CENTER_TEMP].agg(['mean', 'std', 'min', 'max']).to_dict(),
					'by_healing_status': df.groupby('Healing_Status')[CN.CENTER_TEMP].mean().to_dict(),
					'temporal_trend': df.groupby(CN.VISIT_NUMBER)[CN.CENTER_TEMP].mean().to_dict()
				}
			}

			# Add edge and peri-wound temperatures if available
			if all(col in df.columns for col in [CN.EDGE_TEMP, CN.PERI_TEMP]):
				stats_data['sensor_data']['temperature'].update({
					'edge_temp': {
						'overall': df[CN.EDGE_TEMP].agg(['mean', 'std', 'min', 'max']).to_dict(),
						'by_healing_status': df.groupby('Healing_Status')[CN.EDGE_TEMP].mean().to_dict()
					},
					'peri_temp': {
						'overall': df[CN.PERI_TEMP].agg(['mean', 'std', 'min', 'max']).to_dict(),
						'by_healing_status': df.groupby('Healing_Status')[CN.PERI_TEMP].mean().to_dict()
					},
					'gradients': {
						'center_to_edge': (df[CN.CENTER_TEMP] - df[CN.EDGE_TEMP]).agg(['mean', 'std']).to_dict(),
						'center_to_peri': (df[CN.CENTER_TEMP] - df[CN.PERI_TEMP]).agg(['mean', 'std']).to_dict(),
						'by_healing_status': df.groupby('Healing_Status').apply(
							lambda x: {
								'center_to_edge': (x[CN.CENTER_TEMP] - x[CN.EDGE_TEMP]).mean(),
								'center_to_peri': (x[CN.CENTER_TEMP] - x[CN.PERI_TEMP]).mean()
							}
						).to_dict()
					}
				})

		# Impedance Analysis
		if CN.HIGHEST_FREQ_ABSOLUTE in df.columns:
			stats_data['sensor_data']['impedance'] = {
				'magnitude': {
					'overall': df[CN.HIGHEST_FREQ_ABSOLUTE].agg(['mean', 'std', 'min', 'max']).to_dict(),
					'by_healing_status': df.groupby('Healing_Status')[CN.HIGHEST_FREQ_ABSOLUTE].mean().to_dict(),
					'temporal_trend': df.groupby(CN.VISIT_NUMBER)[CN.HIGHEST_FREQ_ABSOLUTE].mean().to_dict()
				}
			}

			# Add complex impedance components if available
			if all(col in df.columns for col in [CN.HIGHEST_FREQ_REAL, CN.HIGHEST_FREQ_IMAGINARY]):
				stats_data['sensor_data']['impedance'].update({
					'complex_components': {
						'real': {
							'overall': df[CN.HIGHEST_FREQ_REAL].agg(['mean', 'std', 'min', 'max']).to_dict(),
							'by_healing_status': df.groupby('Healing_Status')[CN.HIGHEST_FREQ_REAL].mean().to_dict()
						},
						'imaginary': {
							'overall': df[CN.HIGHEST_FREQ_IMAGINARY].agg(['mean', 'std', 'min', 'max']).to_dict(),
							'by_healing_status': df.groupby('Healing_Status')[CN.HIGHEST_FREQ_IMAGINARY].mean().to_dict()
						}
					}
				})

		# Oxygenation Analysis
		if CN.OXYGENATION in df.columns:
			stats_data['sensor_data']['oxygenation'] = {
				'oxygenation': {
					'overall'                 : df[CN.OXYGENATION].agg(['mean', 'std', 'min', 'max']).to_dict(),
					'by_healing_status'       : df.groupby('Healing_Status')[CN.OXYGENATION].mean().to_dict(),
					'temporal_trend'          : df.groupby(CN.VISIT_NUMBER)[CN.OXYGENATION].mean().to_dict(),
					'correlation_with_healing': df[CN.OXYGENATION].corr(df[CN.HEALING_RATE]),
					'distribution_quartiles'  : pd.qcut(df[CN.OXYGENATION], q=4).value_counts().to_dict()
				}
			}

			# Add hemoglobin measurements if available
			for hb_type, col in {'hemoglobin': CN.HEMOGLOBIN,
								'oxyhemoglobin': CN.OXYHEMOGLOBIN,
								'deoxyhemoglobin': CN.DEOXYHEMOGLOBIN}.items():
				if col in df.columns:
					stats_data['sensor_data']['oxygenation'][hb_type] = {
						'overall'                 : df[col].agg(['mean', 'std', 'min', 'max']).to_dict(),
						'by_healing_status'       : df.groupby('Healing_Status')[col].mean().to_dict(),
						'temporal_trend'          : df.groupby(CN.VISIT_NUMBER)[col].mean().to_dict(),
						'correlation_with_healing': df[col].corr(df[CN.HEALING_RATE])
					}

		return stats_data

	def get_processed_data(self) -> pd.DataFrame:
		"""
		Process and transform the loaded wound data to prepare it for analysis.

		This method performs several data processing steps including:
		- Cleaning column names
		- Filtering out skipped visits
		- Extracting visit number information
		- Converting and formatting dates
		- Calculating days since first visit
		- Handling wound type categorization
		- Calculating wound area when not present
		- Converting numeric columns
		- Creating derived features (temperature gradients, BMI categories)
		- Calculating healing rates and metrics for each patient
		- Estimating days to heal based on healing trajectory

		Returns:
			pd.DataFrame: Processed dataframe with additional calculated columns including
						healing rates, temperature gradients, and estimated days to heal.

		Raises:
			ValueError: If no data has been loaded prior to calling this method.

		Note:
			This method does not modify the original dataframe but returns a processed copy.
		"""

		if self.df is None:
			raise ValueError("No data available. Please load data first.")

		# Create a copy to avoid modifying original data
		df = self.df.copy()

		# Clean column names
		df.columns = df.columns.str.strip()

		# Filter out skipped visits
		df = df[df[self.CN.SKIPPED_VISIT] != 'Yes']

		# Extract visit number from Event Name
		df['Visit Number'] = df[self.CN.EVENT_NAME].str.extract(r'Visit (\d+)').fillna(1).astype(int)

		# Convert and format dates
		df[self.CN.VISIT_DATE] = pd.to_datetime(df[self.CN.VISIT_DATE])

		# Calculate days since first visit for each patient
		df[self.CN.DAYS_SINCE_FIRST_VISIT] = df.groupby(self.CN.RECORD_ID)[self.CN.VISIT_DATE].transform(
			lambda x: (x - x.min()).dt.days
		)

		# Handle Wound Type categorization
		if self.CN.WOUND_TYPE in df.columns:

			# Convert to string type first to handle any existing categorical
			df[self.CN.WOUND_TYPE] = df[self.CN.WOUND_TYPE].astype(str)

			# Replace NaN with 'Unknown'
			df[self.CN.WOUND_TYPE] = df[self.CN.WOUND_TYPE].replace('nan', 'Unknown')

			# Get unique categories including 'Unknown'
			categories = sorted(df[self.CN.WOUND_TYPE].unique())

			# Now create categorical with all possible categories
			df[self.CN.WOUND_TYPE] = pd.Categorical(df[self.CN.WOUND_TYPE], categories=categories)

		# Calculate wound area if not present
		if self.CN.WOUND_AREA not in df.columns and all(col in df.columns for col in [self.CN.LENGTH, self.CN.WIDTH]):
			df[self.CN.WOUND_AREA] = df[self.CN.LENGTH] * df[self.CN.WIDTH]

		# Convert numeric columns
		numeric_columns = [self.CN.LENGTH, self.CN.WIDTH, self.CN.HEALING_RATE, self.CN.OXYGENATION]
		for col in numeric_columns:
			if col in df.columns:
				df[col] = pd.to_numeric(df[col].astype(str).str.replace('%', ''), errors='coerce')

		# Create derived features
		# Temperature gradients
		if all(col in df.columns for col in [self.CN.CENTER_TEMP, self.CN.EDGE_TEMP, self.CN.PERI_TEMP]):
			df[self.CN.CENTER_EDGE_GRADIENT] = df[self.CN.CENTER_TEMP] - df[self.CN.EDGE_TEMP]
			df[self.CN.EDGE_PERI_GRADIENT]   = df[self.CN.EDGE_TEMP]   - df[self.CN.PERI_TEMP]
			df[self.CN.TOTAL_GRADIENT]       = df[self.CN.CENTER_TEMP] - df[self.CN.PERI_TEMP]

		# BMI categories
		if self.CN.BMI in df.columns:
			df[self.CN.BMI_CATEGORY] = pd.cut(
				df[self.CN.BMI],
				bins=[0, 18.5, 25, 30, 35, 100],
				labels=['Underweight', 'Normal', 'Overweight', 'Obese I', 'Obese II-III']
			)

		# Calculate healing rates
		MAX_TREATMENT_DAYS = 730  # 2 years in days
		MIN_WOUND_AREA = 0

		def calculate_patient_healing_metrics(patient_data: pd.DataFrame) -> tuple:
			"""
			Calculate healing metrics for a patient based on wound area measurements across visits.

			This function analyzes a patient's wound data over multiple visits to determine
			healing rates, improvement status, and estimated days to complete healing.

			Parameters:
			----------
			patient_data : pd.DataFrame
				DataFrame containing the patient's wound measurements across visits.
				Must contain columns for 'Visit Number' and wound area.

			Returns:
			-------
			tuple
				A tuple containing three elements:
				- healing_rates (list): List of healing rates (percentage) between consecutive visits.
				- is_improving (bool): Flag indicating if the wound is showing improvement (True) or not (False).
				- estimated_days (float or np.nan): Estimated total days until complete healing based on the current healing rate. Returns np.nan if estimation is not possible or unreliable.

			Notes:
			-----
			- Healing rate is calculated as: (previous_area - current_area) / previous_area * 100
			- If patient has fewer than 2 visits, returns ([0.0], False, np.nan)
			- The estimation uses the average healing rate of positive rates only
			- Estimation returns np.nan if:
			  * The wound is not improving
			  * Current area is below MIN_WOUND_AREA
			  * Healing rate is negative or zero
			  * Estimated days exceeds MAX_TREATMENT_DAYS
			"""

			if len(patient_data) < 2:
				return [0.0], False, np.nan

			healing_rates = []
			for i, row in patient_data.iterrows():
				if row[self.CN.VISIT_NUMBER] == 1 or len(patient_data[patient_data[self.CN.VISIT_NUMBER] < row[self.CN.VISIT_NUMBER]]) == 0:
					healing_rates.append(0)
				else:
					prev_visits = patient_data[patient_data[self.CN.VISIT_NUMBER] < row[self.CN.VISIT_NUMBER]]
					prev_visit = prev_visits[prev_visits[self.CN.VISIT_NUMBER] == prev_visits[self.CN.VISIT_NUMBER].max()]

					if len(prev_visit) > 0 and self.CN.WOUND_AREA in patient_data.columns:
						prev_area = prev_visit[self.CN.WOUND_AREA].values[0]
						curr_area = row[self.CN.WOUND_AREA]

						if prev_area > 0 and not pd.isna(prev_area) and not pd.isna(curr_area):
							healing_rate = (prev_area - curr_area) / prev_area * 100
							healing_rates.append(healing_rate)
						else:
							healing_rates.append(0)
					else:
						healing_rates.append(0)

			valid_rates = [rate for rate in healing_rates if rate > 0]
			avg_healing_rate = np.mean(valid_rates) if valid_rates else 0
			is_improving = avg_healing_rate > 0

			estimated_days = np.nan
			if is_improving and len(patient_data) > 0:
				last_visit = patient_data.iloc[-1]
				current_area = last_visit[self.CN.WOUND_AREA]

				if current_area > MIN_WOUND_AREA and avg_healing_rate > 0:
					daily_healing_rate = (avg_healing_rate / 100) * current_area
					if daily_healing_rate > 0:
						days_to_heal = current_area / daily_healing_rate
						total_days = last_visit[self.CN.DAYS_SINCE_FIRST_VISIT] + days_to_heal
						if 0 < total_days < MAX_TREATMENT_DAYS:
							estimated_days = float(total_days)

			return healing_rates, is_improving, estimated_days

		# Process each patient's data
		for patient_id in df[self.CN.RECORD_ID].unique():
			patient_data = df[df[self.CN.RECORD_ID] == patient_id].sort_values(self.CN.DAYS_SINCE_FIRST_VISIT)
			healing_rates, is_improving, estimated_days = calculate_patient_healing_metrics(patient_data)

			# Update patient records with healing rates
			for i, (idx, row) in enumerate(patient_data.iterrows()):
				if i < len(healing_rates):
					df.loc[idx, self.CN.HEALING_RATE] = healing_rates[i]

			# Update the last visit with overall improvement status
			df.loc[patient_data.iloc[-1].name, self.CN.OVERALL_IMPROVEMENT] = 'Yes' if is_improving else 'No'

			if not np.isnan(estimated_days):
				df.loc[patient_data.index, self.CN.ESTIMATED_DAYS_TO_HEAL] = estimated_days

		# Calculate and store average healing rates
		df[self.CN.AVERAGE_HEALING_RATE] = df.groupby(self.CN.RECORD_ID)[self.CN.HEALING_RATE].transform('mean')

		# Ensure estimated days column exists
		if self.CN.ESTIMATED_DAYS_TO_HEAL not in df.columns:
			df[self.CN.ESTIMATED_DAYS_TO_HEAL] = pd.Series(np.nan, index=df.index, dtype=float)

		return df

	def _extract_patient_metadata(self, patient_data) -> Dict[str, Any]:
		"""
		Extracts and formats patient metadata from the provided patient data.

		This method compiles demographic information, lifestyle factors, and medical history
		into a structured dictionary format. It handles missing values by setting them to None
		and processes medical history from both standard columns and free text fields.

		Parameters
		----------
		patient_data : pandas.Series or dict
			A row of patient data containing demographic information, lifestyle factors, and medical history.

		Returns
		-------
		Dict
			A dictionary containing the patient's metadata with the following keys:
			- Basic demographics (age, sex, race, ethnicity)
			- Physical measurements (weight, height, bmi)
			- Study information (study_cohort)
			- Lifestyle factors (smoking_status, packs_per_day, years_smoking, alcohol_use, alcohol_frequency)
			- Medical history (as a nested dictionary)
			- Diabetes information (status, hemoglobin_a1c, a1c_available)

		Notes
		-----
		The method uses column name mappings from the self.columns object to
		access the appropriate fields in the patient data.
		"""

		def safe_value(key):
			value = patient_data.get(key)
			return None if pd.isna(value) else value

		# List of basic metadata columns
		basic_keys = [
				self.CN.AGE,
				self.CN.SEX,
				self.CN.RACE,
				self.CN.ETHNICITY,
				self.CN.WEIGHT,
				self.CN.HEIGHT,
				self.CN.BMI,
				self.CN.STUDY_COHORT,
				self.CN.SMOKING_STATUS,
				self.CN.PACKS_PER_DAY,
				self.CN.YEARS_SMOKED,
				self.CN.ALCOHOL_STATUS,
				self.CN.ALCOHOL_DRINKS
		]
		metadata = {key: safe_value(key) for key in basic_keys}


		# Medical history from individual columns
		medical_conditions = [
			self.CN.RESPIRATORY, self.CN.CARDIOVASCULAR, self.CN.GASTROINTESTINAL, self.CN.MUSCULOSKELETAL,
			self.CN.ENDOCRINE_METABOLIC, self.CN.HEMATOPOIETIC, self.CN.HEPATIC_RENAL, self.CN.NEUROLOGIC, self.CN.IMMUNE
		]

		# Get medical history from standard columns
		metadata['medical_history'] = {
			condition: patient_data[condition]
			for condition in medical_conditions if not pd.isna(patient_data.get(condition))
		}

		# Check additional medical history from free text field
		other_history = patient_data.get(self.CN.MEDICAL_HISTORY)
		if not pd.isna(other_history):
			existing_conditions = set(medical_conditions)
			other_conditions = [cond.strip() for cond in str(other_history).split(',')]
			other_conditions = [cond for cond in other_conditions if cond and cond not in existing_conditions]
			if other_conditions:
				metadata['medical_history']['other'] = ', '.join(other_conditions)

		# Diabetes information
		metadata['diabetes'] = {
			self.CN.DIABETES        : patient_data.get(self.CN.DIABETES),
			self.CN.A1C             : patient_data.get(self.CN.A1C),
			self.CN.A1C_AVAILABLE   : patient_data.get(self.CN.A1C_AVAILABLE)
		}

		return metadata

	def _get_wound_info(self, visit_data) -> Dict:
		"""
		Extract and structure detailed wound information from a single visit record.

		This method processes raw visit data to extract wound characteristics and clinical assessment
		information into a structured dictionary format. It handles missing data by converting NaN values
		to None.

		Args:
			visit_data (Dict): A dictionary containing wound data from a single visit

		Returns:
			Dict: Structured wound information with the following keys:
				- location: Anatomical location of the wound
				- type: Classification of wound type
				- current_care: Current wound care regimen
				- clinical_events: Notable clinical events
				- undermining: Dictionary containing undermining presence, location and tunneling details
				- infection: Dictionary with infection status and WiFi classification
				- granulation: Dictionary with tissue coverage and quality metrics
				- necrosis: Necrotic tissue assessment
				- exudate: Dictionary containing volume, viscosity and type of wound drainage

		Raises:
			Exception: Logs a warning and returns an empty dictionary if data processing fails
		"""
		def clean_field(field):
			return visit_data.get(field) if not pd.isna(visit_data.get(field)) else None

		present = clean_field(self.CN.UNDERMINING)

		wound_info = {
			'location'       : clean_field(self.CN.WOUND_LOCATION),
			'type'           : clean_field(self.CN.WOUND_TYPE),
			'current_care'   : clean_field(self.CN.CURRENT_WOUND_CARE),
			'clinical_events': clean_field(self.CN.CLINICAL_EVENTS),
			'undermining': {
				'present'  : None if present is None else present == 'Yes',
				'location' : clean_field(self.CN.UNDERMINING_LOCATION),
				'tunneling': clean_field(self.CN.TUNNELING_LOCATION)
			},
			'infection': {
				'status'             : clean_field(self.CN.INFECTION),
				'wifi_classification': clean_field(self.CN.WIFI_CLASSIFICATION)
			},
			'granulation': {
				'coverage': clean_field(self.CN.GRANULATION),
				'quality' : clean_field(self.CN.GRANULATION_QUALITY)
			},
			'necrosis': clean_field(self.CN.NECROSIS),
			'exudate': {
				'volume'   : clean_field(self.CN.EXUDATE_VOLUME),
				'viscosity': clean_field(self.CN.EXUDATE_VISCOSITY),
				'type'     : clean_field(self.CN.EXUDATE_TYPE)
			}
		}

		return wound_info


	def _process_visit_data(self, visit_csv_row: pd.Series, record_id: int) -> Optional[VisitsDataType]:
		"""
			Process the data from a single patient visit and extract relevant information.

			This method extracts and processes various measurements taken during a patient visit, including
			wound measurements, temperature readings, oxygenation data, and impedance values. It handles
			missing data gracefully by converting NaN values to None.

			Args:
				visit: A dictionary-like object containing the raw visit data
				record_id (int): The unique identifier for the patient record

			Returns:
				Optional[Dict]: A structured dictionary containing the processed visit data, or None if
				the visit date is missing. The dictionary includes:
					- visit_date: The formatted date of the visit
					- wound_measurements: Dict containing length, width, depth, and area
					- sensor_data: Dict containing:
						- oxygenation: Overall oxygenation value
						- temperature: Dict with center, edge, and peri readings
						- impedance: Dict with high, center, and low frequency measurements
						- hemoglobin: Various hemoglobin measurements

			Notes:
				The impedance data is either extracted from Excel sweep files if available,
				or from the visit parameters directly as a fallback.

				This method uses a caching mechanism to avoid reprocessing impedance data
				for the same patient visit, which improves performance significantly.
		"""

		# Extract visit date
		visit_date = pd.to_datetime(visit_csv_row[self.CN.VISIT_DATE]).strftime('%m-%d-%Y') if not pd.isna(visit_csv_row.get(self.CN.VISIT_DATE)) else None

		if not visit_date:
			logger.warning("Missing visit date")
			return None

		# Check if we have cached impedance data for this visit
		visit_number = visit_csv_row.get(self.CN.VISIT_NUMBER)

		# logger.info(f"Processing impedance data for patient {record_id}, visit {visit_number}")
		impedance_data = self.impedance_analyzer.get_structured_three_freq_high_center_low_impedance_data_for_view(visit_csv_row=visit_csv_row, record_id=record_id, visit_date=visit_date, CN=self.CN)

		wound_measurements = {
			'length': get_float(visit_csv_row, self.CN.LENGTH),
			'width' : get_float(visit_csv_row, self.CN.WIDTH),
			'depth' : get_float(visit_csv_row, self.CN.DEPTH),
			'area'  : get_float(visit_csv_row, self.CN.WOUND_AREA)
		}

		temperature_readings = {
			'center': get_float(visit_csv_row, self.CN.CENTER_TEMP),
			'edge'  : get_float(visit_csv_row, self.CN.EDGE_TEMP),
			'peri'  : get_float(visit_csv_row, self.CN.PERI_TEMP)
		}

		hemoglobin_types = {
			'hemoglobin'     : self.CN.HEMOGLOBIN,
			'oxyhemoglobin'  : self.CN.OXYHEMOGLOBIN,
			'deoxyhemoglobin': self.CN.DEOXYHEMOGLOBIN
		}

		return {
			self.CN.VISIT_DATE: visit_date,
			'wound_measurements': wound_measurements,
			'sensor_data': {
				'oxygenation': get_float(visit_csv_row, self.CN.OXYGENATION),
				'temperature': temperature_readings,
				'impedance'  : impedance_data,
				**{key: get_float(visit_csv_row, value) for key, value in hemoglobin_types.items()}
			}
		}


	def get_patient_dataframe(self, record_id: int) -> pd.DataFrame:
		"""
		Get the patient dataframe for a given record ID.
		"""
		return self.df[self.df[self.CN.RECORD_ID] == record_id].sort_values(self.CN.VISIT_DATE)


	@staticmethod
	def get_visit_date_tag(visits:List[VisitsDataType]) -> str:
		return [k for k in visits[0].keys() if k not in ['wound_measurements', 'sensor_data']][0]


@dataclass
class DataManager:
	"""Handles data loading, processing and manipulation."""
	df     : pd.DataFrame = field(default_factory=lambda: None)
	CN     : DColumns     = field(default_factory=lambda: None)

	@staticmethod
	def load_data(csv_dataset_path):
		"""Load and preprocess the wound healing data from an uploaded CSV file. Returns None if no file is provided."""

		if csv_dataset_path is None:
			return None

		df = pd.read_csv(csv_dataset_path)
		df = DataManager._preprocess_data(df)
		return df

	@staticmethod
	def _preprocess_data(df: pd.DataFrame) -> pd.DataFrame:
		"""
		Preprocesses the wound data DataFrame.

		This function performs several preprocessing steps on the input DataFrame, including:
		1. Normalizing column names
		2. Filtering out skipped visits
		3. Extracting and converting visit numbers
		4. Formatting visit dates
		5. Converting wound type to categorical data
		6. Calculating wound area from dimensions if not present
		7. Creating derived features and healing rates

		Args:
			df (pd.DataFrame): The raw wound data DataFrame to preprocess.

		Returns:
			pd.DataFrame: The preprocessed wound data.
		"""

		df = df.copy()

		# Get column names from schema
		CN = DColumns(df=df)

		df.columns = df.columns.str.strip()

		# Filter out skipped visits
		df = df[df[CN.SKIPPED_VISIT] != 'Yes']

		# Fill missing Visit Number with 1 before converting to int
		df[CN.VISIT_NUMBER] = df[CN.EVENT_NAME].str.extract(r'Visit (\d+)').fillna(1).astype(int)

		df[CN.VISIT_DATE] = pd.to_datetime(df[CN.VISIT_DATE]).dt.strftime('%m-%d-%Y')

		# Convert Wound Type to categorical with specified categories
		df[CN.WOUND_TYPE] = pd.Categorical(df[CN.WOUND_TYPE].fillna('Unknown'), categories=df[CN.WOUND_TYPE].dropna().unique())

		# 2. Calculate wound area if not present but dimensions are available
		if CN.WOUND_AREA not in df.columns and all(col in df.columns for col in [CN.LENGTH, CN.WIDTH]):
			df[CN.WOUND_AREA] = df[CN.LENGTH] * df[CN.WIDTH]

		df = DataManager._create_derived_features(df)
		df = DataManager._calculate_healing_rates(df)
		return df

	@staticmethod
	def _calculate_healing_rates(df: pd.DataFrame) -> pd.DataFrame:
		"""Calculate healing rates and related metrics for each patient's wound data.

		This function processes wound measurement data to calculate:
		- Per-visit healing rates (percentage change in wound area)
		- Overall improvement status (Yes/No)
		- Average healing rate across all visits
		- Estimated days to complete healing

		The calculations are performed on a per-patient basis and consider the
		sequential nature of wound healing across multiple visits.

		Args:
			df (pd.DataFrame): DataFrame containing wound measurement data with columns defined in the DataColumns schema.

			pd.DataFrame: The input DataFrame with additional columns for healing metrics:
							- healing_rate: Per-visit healing rate
							- overall_improvement: Yes/No indicator of healing progress
							- average_healing_rate: Mean healing rate across visits
							- estimated_days_to_heal: Projected time to complete healing

		Notes:
			- Healing rate is calculated as percentage decrease in wound area between visits
			- A positive healing rate indicates improvement (wound area reduction)
			- Estimated healing time is based on the latest wound area and average healing rate
			- Calculations handle edge cases like single visits and invalid measurements
		"""

		# Get column names from schema
		CN = DColumns(df=df)

		# Constants
		MAX_TREATMENT_DAYS = 730  # 2 years in days
		MIN_WOUND_AREA = 0

		def calculate_patient_healing_metrics(patient_data: pd.DataFrame) -> tuple[list, bool, float]:
			"""Calculate healing rate and estimated days for a patient.

			Returns:
				tuple: (healing_rates, is_improving, estimated_days_to_heal)
			"""
			if len(patient_data) < 2:
				return [0.0], False, np.nan

			# 3. Calculate healing rate (% change in wound area per visit)
			healing_rates = []
			for i, row in patient_data.iterrows():
				# If first visit or no previous visits exist, healing rate is 0
				if row[CN.VISIT_NUMBER] == 1:
					healing_rates.append(0)
					continue

				# Get the most recent previous visit using idxmax() instead of filtering twice
				prev_visits = patient_data[patient_data[CN.VISIT_NUMBER] < row[CN.VISIT_NUMBER]]
				if prev_visits.empty:
					healing_rates.append(0)
					continue

				prev_visit_idx = prev_visits[CN.VISIT_NUMBER].idxmax()
				prev_area = patient_data.loc[prev_visit_idx, CN.WOUND_AREA]
				curr_area = row[CN.WOUND_AREA]

				# Calculate healing rate if values are valid
				if pd.notna(prev_area) and pd.notna(curr_area) and prev_area > 0:
					healing_rate = (prev_area - curr_area) / prev_area * 100
					healing_rates.append(healing_rate)
				else:
					healing_rates.append(0)

			# Calculate average healing rate and determine if improving
			valid_rates = [rate for rate in healing_rates if rate > 0]
			avg_healing_rate = np.mean(valid_rates) if valid_rates else 0
			is_improving = avg_healing_rate > 0

			# Calculate estimated days to heal based on the latest wound area and average healing rate
			estimated_days = np.nan
			if is_improving and len(patient_data) > 0:
				last_visit = patient_data.iloc[-1]
				current_area = last_visit[CN.WOUND_AREA]

				if current_area > MIN_WOUND_AREA and avg_healing_rate > 0:
					# Convert percentage rate to area change per day
					daily_healing_rate = (avg_healing_rate / 100) * current_area
					if daily_healing_rate > 0:
						days_to_heal = current_area / daily_healing_rate
						total_days = last_visit[CN.DAYS_SINCE_FIRST_VISIT] + days_to_heal
						if 0 < total_days < MAX_TREATMENT_DAYS:
							estimated_days = float(total_days)

			return healing_rates, is_improving, estimated_days

		# Process each patient's data
		for patient_id in df[CN.RECORD_ID].unique():
			patient_data = df[df[CN.RECORD_ID] == patient_id].sort_values(CN.DAYS_SINCE_FIRST_VISIT)

			healing_rates, is_improving, estimated_days = calculate_patient_healing_metrics(patient_data)

			# Update patient records with healing rates
			for i, (idx, row) in enumerate(patient_data.iterrows()):
				if i < len(healing_rates):
					df.loc[idx, CN.HEALING_RATE] = healing_rates[i]

			# Update the last visit with overall improvement status
			df.loc[patient_data.iloc[-1].name, CN.OVERALL_IMPROVEMENT] = 'Yes' if is_improving else 'No'

			if not np.isnan(estimated_days):
				df.loc[patient_data.index, CN.ESTIMATED_DAYS_TO_HEAL] = estimated_days

		# Calculate and store average healing rates
		df[CN.AVERAGE_HEALING_RATE] = df.groupby(CN.RECORD_ID)[CN.HEALING_RATE].transform('mean')

		# Ensure estimated days column exists
		if CN.ESTIMATED_DAYS_TO_HEAL not in df.columns:
			df[CN.ESTIMATED_DAYS_TO_HEAL] = pd.Series(np.nan, index=df.index, dtype=float)

		return df

	@staticmethod
	def _create_derived_features(df: pd.DataFrame) -> pd.DataFrame:
		"""
		Creates derived features from raw dataframe columns.

		This function adds new calculated columns to the input DataFrame based on existing data:
		- Temperature gradients (center-edge, edge-perimeter, total gradients)
		- BMI categorization using standard health ranges
		- Days since first visit for each patient
		- Initializes healing metric columns (healing rate, estimated days to heal, overall improvement)

		Parameters
		----------
		df : pd.DataFrame
			Input DataFrame containing wound assessment data with columns specified in DataColumns schema

		Returns
		-------
		pd.DataFrame
			DataFrame with original columns plus newly derived features
		"""

		# Make a copy to avoid SettingWithCopyWarning
		df = df.copy()

		# Get column names from schema
		CN = DColumns(df=df)


		# Temperature gradients
		if all(col in df.columns for col in [CN.CENTER_TEMP, CN.EDGE_TEMP, CN.PERI_TEMP]):
			df[CN.CENTER_EDGE_GRADIENT] = df[CN.CENTER_TEMP] - df[CN.EDGE_TEMP]
			df[CN.EDGE_PERI_GRADIENT]   = df[CN.EDGE_TEMP]   - df[CN.PERI_TEMP]
			df[CN.TOTAL_GRADIENT]       = df[CN.CENTER_TEMP] - df[CN.PERI_TEMP]

		# BMI categories
		if CN.BMI in df.columns:
			df[CN.BMI_CATEGORY] = pd.cut(
				df[CN.BMI],
				bins=[0, 18.5, 25, 30, 35, 100],
				labels=['Underweight', 'Normal', 'Overweight', 'Obese I', 'Obese II-III']
			)

		if df is not None and not df.empty:
			# Convert Visit date to datetime if not already
			df[CN.VISIT_DATE] = pd.to_datetime(df[CN.VISIT_DATE])

			# Calculate days since first visit for each patient
			df[CN.DAYS_SINCE_FIRST_VISIT] = df.groupby(CN.RECORD_ID)[CN.VISIT_DATE].transform(
				lambda x: (x - x.min()).dt.days
			)

			# Initialize columns with explicit dtypes
			df.loc[:, CN.HEALING_RATE]           = pd.Series(0.0, index=df.index, dtype=float)
			df.loc[:, CN.ESTIMATED_DAYS_TO_HEAL] = pd.Series(np.nan, index=df.index, dtype=float)
			df.loc[:, CN.OVERALL_IMPROVEMENT]    = pd.Series(np.nan, index=df.index, dtype=str)

		return df


	@staticmethod
	def create_and_save_report(patient_metadata: dict, analysis_results: str, report_path: str=None, prompt: dict=None) -> str:
		"""
			Create a report document from analysis results and save it.

			This function creates a Word document containing the analysis results for a patient,
			formats it using the WoundDataProcessor utility, and saves it to disk.

			Parameters
			----------
			patient_metadata : dict
				Dictionary containing patient information such as name, ID, and demographic data.
			analysis_results : str
				The analysis results text to be included in the report.
			prompt : dict, optional
				Optional dictionary containing prompt information that was used to generate the analysis.

			Returns
			-------
			str
				The file path of the saved report document.
		"""
		# Create logs directory if report_path is not provided
		if report_path is None:
			log_dir = pathlib.Path(__file__).parent / 'logs'
			log_dir.mkdir(exist_ok=True)
			timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
			report_path = log_dir / f'wound_analysis_{timestamp}.docx'

		# Create and format document
		doc = Document()
		DataManager.format_word_document(doc=doc, analysis_results=analysis_results, patient_metadata=patient_metadata, report_path=report_path)

		return report_path

	@staticmethod
	def download_word_report(st, report_path: str):
		"""
			Creates a download button in a Streamlit app to download a Word document (.docx) report.

			Parameters:
			-----------
			st : streamlit
				The Streamlit module instance used to render UI components.
			report_path : str
				The file path to the Word document (.docx) to be downloaded.

			Returns:
			--------
			None

			Raises:
			-------
			Exception
				If there's an error reading the report file or creating the download button,
				an error message is displayed in the Streamlit app.

			Notes:
			------
			This function reads the specified Word document as binary data and creates
			a download button in the Streamlit interface.
		"""

		try:
			with open(report_path, 'rb') as f:
				bytes_data = f.read()
				st.download_button(
					label="Download Full Report (DOCX)",
					data=bytes_data,
					file_name=os.path.basename(report_path),
					mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
				)
		except Exception as e:
			st.error(f"Error preparing report download: {str(e)}")

	@staticmethod
	def format_word_document(doc: Document, analysis_results: str, patient_metadata: dict=None, report_path: str = None):
		"""
			Formats and saves a Microsoft Word document with wound care analysis results.

			This function creates a structured report document with patient information and
			analysis results, applying proper formatting to different sections of content.

			Parameters
			----------
			doc : Document
				A python-docx Document object to be formatted.
			analysis_results : str
				The wound care analysis results as a string, potentially containing markdown-style
				formatting like '**' for bold text and bullet points.
			patient_metadata : dict, optional
				Dictionary containing patient information with keys such as 'age', 'sex', 'bmi',
				and nested 'diabetes' information. If None, patient information section is omitted.
			report_path : str, optional
				Path where the document should be saved. If None, a default path is generated
				in a 'logs' directory with a timestamp.

			Returns
			-------
			str
				String representation of the path where the document was saved.

			Notes
			-----
			The function handles markdown-style formatting in the analysis_results:
			- Text surrounded by '**' is converted to bold
			- Lines starting with '- ' or '* ' are converted to bullet points
		"""

		# Add title
		title = doc.add_heading('Wound Care Analysis Report', 0)
		title.alignment = WD_ALIGN_PARAGRAPH.CENTER

		if patient_metadata is not None:
			# Add patient information section
			doc.add_heading('Patient Information', level=1)
			patient_info = doc.add_paragraph()
			patient_info.add_run('Patient Demographics:\n').bold = True
			patient_info.add_run(f"Age: {patient_metadata.get('age', 'Unknown')} years\n")
			patient_info.add_run(f"Sex: {patient_metadata.get('sex', 'Unknown')}\n")
			patient_info.add_run(f"BMI: {patient_metadata.get('bmi', 'Unknown')}\n")

			# Add diabetes information
			diabetes_info = doc.add_paragraph()
			diabetes_info.add_run('Diabetes Status:\n').bold = True
			if 'diabetes' in patient_metadata:
				diabetes_info.add_run(f"Type: {patient_metadata['diabetes'].get('status', 'Unknown')}\n")
				diabetes_info.add_run(f"HbA1c: {patient_metadata['diabetes'].get('hemoglobin_a1c', 'Unknown')}%\n")

		# Add analysis section
		doc.add_heading('Analysis Results', level=1)

		# Split analysis into sections and format them
		sections = analysis_results.split('\n\n')
		for section in sections:
			if section.strip():
				if '**' in section:  # Handle markdown-style headers
					# Convert markdown headers to proper formatting
					section = section.replace('**', '')
					p = doc.add_paragraph()
					p.add_run(section.strip()).bold = True
				else:
					# Handle bullet points
					if section.strip().startswith('- ') or section.strip().startswith('* '):
						p = doc.add_paragraph(section.strip()[2:], style='List Bullet')
					else:
						p = doc.add_paragraph(section.strip())

		# Add footer with timestamp
		doc.add_paragraph(f"\nReport generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

		# Save the document
		if report_path is not None:
			doc.save(report_path)


class ImpedanceExcelProcessor:
	"""
		Processes and manages impedance data from Excel files containing frequency sweep measurements.

		This class handles the loading, parsing, and caching of impedance frequency sweep data
		from Excel files. It extracts key frequency data points (lowest, center, highest frequencies)
		for each patient visit and provides methods to retrieve specific visit data.

		The processor maintains a cache to improve performance when repeatedly accessing
		the same data files.

		Attributes:
			cache : dict
				A dictionary that stores processed data to avoid redundant file operations.
				Data is organized by cache keys (based on record ID and file path).

			CN : DColumns
			Column names for the data processor.
	"""
	def __init__(self, impedance_freq_sweep_path: pathlib.Path):
		self.impedance_freq_sweep_path = impedance_freq_sweep_path

		self.cache_processed_excel_files = {}
		self._process_all_excel_files()


	def get_cache_key_excel_file(self, record_id):
		return f"excel_file_{record_id}"


	def _parse_date(self, date_cell_value, sheet_name: str) -> Optional[str]:
		try:
			# Convert to string and extract date part
			date_str = str(date_cell_value)
			date_part = date_str.split(' ')[0] if ' ' in date_str else date_str

			logger.debug(f"Parsing date from sheet {sheet_name}: {date_part}")

			# Try different date formats
			date_formats = [
				"%Y-%d-%m",  # 2024-30-08
				"%d/%m/%Y",  # 30/08/2024
				"%m/%d/%Y",  # 08/30/2024
				"%Y/%m/%d",  # 2024/08/30
				"%d-%m-%Y",  # 30-08-2024
				"%m-%d-%Y",  # 08-30-2024
				"%Y-%m-%d",  # 2024-08-30
			]

			for date_format in date_formats:
				try:
					parsed_date = datetime.strptime(date_part, date_format).strftime('%m-%d-%Y')
					logger.debug(f"Successfully parsed date {date_part} using format {date_format} -> {parsed_date}")
					return parsed_date
				except ValueError:
					continue

			logger.warning(f"Could not parse date '{date_part}' in sheet {sheet_name} with any format")
			return None

		except Exception as e:
			logger.warning(f"Error parsing date: {e}")
			return None


	def _extract_three_frequency_data(self, df_bottom: pd.DataFrame) -> pd.DataFrame:
		"""
			Extract frequency data at key frequency points from the dataframe.

			This method identifies the lowest, highest, and center frequency points
			from the impedance data and extracts corresponding electrical measurements.
			The center frequency is determined as the frequency with the maximum negative phase.

			Parameters
			----------
			freq_data : pd.DataFrame
				DataFrame containing the frequency data with columns:
				'frequency', 'Z', 'Z_prime', 'Z_double_prime', 'neg. Phase / °'

			Returns
			-------
			list
				List of dictionaries containing impedance data at the three key frequencies:
				- lowest frequency
				- center frequency
				- highest frequency
		"""
		# Check if dataframe is empty
		if df_bottom.empty:
			return []

		# Find key frequencies
		lowest_freq  = df_bottom[ExcelSheetColumns.FREQ.value].min()
		highest_freq = df_bottom[ExcelSheetColumns.FREQ.value].max()

		# Find center frequency (max negative phase)
		max_phase_idx = df_bottom[ExcelSheetColumns.NEG_PHASE.value].idxmax()
		center_freq   = df_bottom.loc[max_phase_idx, ExcelSheetColumns.FREQ.value]

		# Create data for each frequency
		result = []
		for freq_type, freq in [('lowest_freq', lowest_freq),
								('center_freq', center_freq),
								('highest_freq', highest_freq)]:

			# Find row for current frequency
			rows = df_bottom[df_bottom[ExcelSheetColumns.FREQ.value] == freq]
			if rows.empty:
				continue

			row = rows.iloc[0]

			result.append({
				'index'                                      : freq_type,
				ExcelSheetColumns.FREQ.value                 : str(freq),
				ExcelSheetColumns.VISIT_DATE_FREQ_SWEEP.value: df_bottom.attrs.get(ExcelSheetColumns.VISIT_DATE_FREQ_SWEEP.value),
				ExcelSheetColumns.ABSOLUTE.value             : row[ExcelSheetColumns.ABSOLUTE.value],
				ExcelSheetColumns.REAL.value                 : row[ExcelSheetColumns.REAL.value],
				ExcelSheetColumns.IMAGINARY.value            : row[ExcelSheetColumns.IMAGINARY.value],
				ExcelSheetColumns.NEG_PHASE.value            : row[ExcelSheetColumns.NEG_PHASE.value]
				})

		return pd.DataFrame(result).set_index('index')


	def _process_sheet(self, excel_file: pathlib.Path, sheet_name: str) -> Optional[pd.DataFrame]:
		"""
			Process a single sheet from an Excel file to extract date and frequency data.

			This method reads the header to extract the date, validates the sheet structure,
			and processes the bottom half of the sheet to extract frequency data.

			Parameters
			----------
			excel_file : pathlib.Path
				Path to the Excel file to process
			sheet_name : str
				Name of the sheet to process within the Excel file

			Returns
			-------
			Optional[pd.DataFrame]
				A DataFrame containing frequency data if successful,
				None if any validation fails or an exception occurs

			Notes
			-----
			The method expects:
			- Date information in cell C1 (index [0,2])
			- Data rows starting from the second row (skiprows=1)
			- Processes only the bottom half of the available data rows
		"""

		# Read header row
		header_df = pd.read_excel(excel_file, sheet_name=sheet_name, nrows=1, header=None)

		# Validate header
		if header_df.shape[1] < 3:
			# logger.warning(f"Sheet {sheet_name} has insufficient columns")
			return None

		# Extract and parse date
		date_cell_value = header_df.iloc[0, 2]
		if pd.isna(date_cell_value):
			logger.warning(f"No date found in sheet {sheet_name}")
			return None

		visit_date_freq_sweep: str = self._parse_date(date_cell_value, sheet_name)
		if not visit_date_freq_sweep:
			return None

		# Process data
		df = pd.read_excel(excel_file, sheet_name=sheet_name, skiprows=1)
		df.columns = df.columns.str.strip()

		# Get bottom half
		half_point = len(df) // 2
		df_bottom = df.iloc[half_point:]

		# date of visit
		df_bottom.attrs[ExcelSheetColumns.VISIT_DATE_FREQ_SWEEP.value] = visit_date_freq_sweep

		return df_bottom


	def _process_excel_file(self, record_id: int) -> None:
		"""
			Process an Excel file containing impedance frequency sweep data for a specific record.

			This method extracts data from each sheet in the Excel file, where each sheet typically
			represents a visit date. The processed data is stored in a dictionary organized by visit date
			and cached using the provided cache key.

			Parameters
			----------
			impedance_freq_sweep_path : pathlib.Path
				Path to the directory containing impedance frequency sweep Excel files
			record_id : int
				The record ID used to locate the specific Excel file
			cache_key : str
				Key to use for storing the processed data in the cache

			Returns
			-------
			None
				The method stores the processed data in the cache under cache_key

			Notes
			-----
			- The Excel file should be named as '{record_id}.xlsx'
			- If the Excel file doesn't exist, a warning is logged and the method returns
			- Exceptions during file processing are caught and logged as errors
		"""

		cache_key = self.get_cache_key_excel_file(record_id)

		if cache_key in self.cache_processed_excel_files:
			return self.cache_processed_excel_files[cache_key]

		# Find the Excel file
		excel_file = self.impedance_freq_sweep_path / f'{record_id}.xlsx'
		if not excel_file.exists():
			logger.warning(f"Failed to find Excel file for record ID: {excel_file}")
			return None

		# Dictionary to store visit data by date
		visit_data_by_date_dict: Dict[str, Dict[str, pd.DataFrame]] = {}

		# Use context manager for proper resource handling
		with pd.ExcelFile(excel_file) as xl:
			for sheet_name in xl.sheet_names:
				# Process each sheet
				df_bottom = self._process_sheet(excel_file=excel_file, sheet_name=sheet_name)

				if df_bottom is not None:
					date = df_bottom.attrs[ExcelSheetColumns.VISIT_DATE_FREQ_SWEEP.value]
					visit_data_by_date_dict[date] = {	'data_three_freqs' : self._extract_three_frequency_data(df_bottom),
														'entire_freq_sweep': df_bottom.copy()}

		# Store in cache
		self.cache_processed_excel_files[cache_key] = visit_data_by_date_dict

		logger.info(f"Processed {len(visit_data_by_date_dict)} visits for record ID: {record_id}")


	def _process_all_excel_files(self) -> None:
		"""
			Processes all Excel files in the specified directory.

			This method iterates through all Excel files in the directory and processes
			each file using the _process_excel_file method. The processed data is stored
			in the cache for future retrieval.
		"""
		# Get sorted list of Excel files
		excel_files = sorted([f for f in self.impedance_freq_sweep_path.iterdir()
							if f.is_file() and f.suffix == '.xlsx' and not f.name.startswith('~')])

		# Only show progress bar if there are files to process
		if excel_files:
			# Create a progress bar
			progress_text = "Loading impedance frequency sweep files..."
			st.write('')
			progress_bar = st.progress(0, text=progress_text)
			total_files = len(excel_files)

			# Process each file in sorted order with progress updates
			for i, file in enumerate(excel_files):
				record_id = int(file.stem)
				self._process_excel_file(record_id=record_id)

				# Update progress bar
				progress = (i + 1) / total_files
				progress_bar.progress(progress, text=f"{progress_text} ({i+1}/{total_files})")

			# Complete the progress bar
			progress_bar.progress(1.0, text="Impedance files loaded successfully!")

			# Add a small delay so users can see the completion message
			import time
			time.sleep(0.5)

			# Clear the progress bar after completion
			progress_bar.empty()
		else:
			# No files to process
			pass


	def get_freq_data_for_visit(self, record_id: int, visit_date: str) -> Optional[Dict[str, pd.DataFrame]]:
		"""
			Retrieves frequency sweep impedance data for a specific visit.

			This method checks the cache for previously processed data before extracting
			visit-specific data from an Excel file. It handles the data retrieval and
			processing workflow.

			Parameters
			----------
			impedance_freq_sweep_path : pathlib.Path
				Path to the Excel file containing impedance frequency sweep data.
			record_id : int
				Unique identifier for the patient record.
			visit_date : str
				Date of the visit for which data is requested.

			Returns
			-------
			Optional[Dict[str, pd.DataFrame]]
				Dictionary containing impedance data for the specified visit date, with 'index'
				set as the index column. Returns None if no data is available for the specified date.

			Note:
				This method uses caching to improve performance. The cache is checked first
				before attempting to process the Excel file again.
		"""
		# Check cache first
		cache_key = self.get_cache_key_excel_file(record_id)

		if cache_key in self.cache_processed_excel_files and visit_date in self.cache_processed_excel_files[cache_key]:
			# logger.debug(f"Using cached frequency data for patient {record_id}, visit date {visit_date}")
			return self.cache_processed_excel_files[cache_key][visit_date]

		return None


@st.cache_resource
class ImpedanceAnalyzer:
	"""Handles advanced bioimpedance analysis and clinical interpretation."""

	def __init__(self, impedance_freq_sweep_path: pathlib.Path):
		self.impedance_freq_sweep_path = impedance_freq_sweep_path
		self._excel_processor = ImpedanceExcelProcessor(impedance_freq_sweep_path)
		self._structured_three_freq_cache = {}  # Cache for processed impedance data


	def get_cache_key_structured_three_freq(self, record_id: int, visit_date: str) -> str:
		return f"{record_id}_{visit_date}"


	def get_structured_three_freq_high_center_low_impedance_data_for_view(self, visit_csv_row: pd.Series, record_id: int, visit_date: str, CN: DColumns) -> Dict:

		# Create a cache key based on record_id and visit date
		cache_key = self.get_cache_key_structured_three_freq(record_id=record_id, visit_date=visit_date)

		# Check if we already have this data in cache
		if cache_key in self._structured_three_freq_cache:
			return self._structured_three_freq_cache[cache_key]

		impedance_data = {
			'highest_freq': {'Z': None, 'resistance': None, 'capacitance': None, 'frequency': None},
			'center_freq' : {'Z': None, 'resistance': None, 'capacitance': None, 'frequency': None},
			'lowest_freq' : {'Z': None, 'resistance': None, 'capacitance': None, 'frequency': None},
			'entire_freq_sweep': None
		}

		def transform_impedance_data(freq_data):
			frequency = get_float(freq_data, ExcelSheetColumns.FREQ.value)

			if freq_data is not None:
				result = {
					'Z'          : freq_data[ExcelSheetColumns.ABSOLUTE.value],
					'resistance' : freq_data[ExcelSheetColumns.REAL.value],
					'capacitance': None if freq_data[ExcelSheetColumns.IMAGINARY.value] is None else 1 / (2 * 3.14 * frequency * freq_data[ExcelSheetColumns.IMAGINARY.value]),
					'frequency'  : frequency
				}
				return result

			return {'Z': None, 'resistance': None, 'capacitance': None, 'frequency': None}


		# Get data for this visit date
		visit_data_for_date: Dict[str, pd.DataFrame] = self._excel_processor.get_freq_data_for_visit(record_id=record_id, visit_date=visit_date)


		if visit_data_for_date is not None:
			df_bottom = visit_data_for_date['entire_freq_sweep']
			data_three_freqs = visit_data_for_date['data_three_freqs'].T
			impedance_data['highest_freq'] = transform_impedance_data(data_three_freqs['highest_freq'].to_dict())
			impedance_data['center_freq']  = transform_impedance_data(data_three_freqs['center_freq'].to_dict())
			impedance_data['lowest_freq']  = transform_impedance_data(data_three_freqs['lowest_freq'].to_dict())
			impedance_data['entire_freq_sweep'] = df_bottom

		else:
			# Get impedance data from visit parameters if no sweep data
			logger.debug("No freq sweep impedance data found")

			high_freq = {
				ExcelSheetColumns.ABSOLUTE.value : get_float(visit_csv_row, CN.HIGHEST_FREQ_ABSOLUTE),
				ExcelSheetColumns.REAL.value     : get_float(visit_csv_row, CN.HIGHEST_FREQ_REAL),
				ExcelSheetColumns.IMAGINARY.value: get_float(visit_csv_row, CN.HIGHEST_FREQ_IMAGINARY),
				ExcelSheetColumns.FREQ.value     : 80000
			}

			impedance_data['highest_freq'] = transform_impedance_data(high_freq)

		# Store the processed data in cache for future use
		self._structured_three_freq_cache[cache_key] = impedance_data

		return impedance_data

	@staticmethod
	def calculate_visit_changes(current_visit: VisitsDataType, previous_visit: VisitsDataType) -> tuple[dict, dict]:
		"""
			Calculate the percentage changes in impedance parameters between two visits and determine clinical significance.

			This function compares impedance data (Z, resistance, capacitance) at different frequencies
			between the current visit and a previous visit. It calculates percentage changes and determines
			if these changes are clinically significant based on predefined thresholds.

			Parameters
			----------
			current_visit : dict
				Dictionary containing the current visit data with sensor_data.impedance measurements
				at low, center, and high frequencies.
			previous_visit : dict
				Dictionary containing the previous visit data with sensor_data.impedance measurements
				at low, center, and high frequencies.

			Returns
			-------
			tuple
				A tuple containing two dictionaries:
				- changes: Dictionary mapping parameter names (e.g., 'resistance_lowest_freq') to their percentage changes between visits.
				- clinically_significant: Dictionary mapping parameter names to boolean values indicating whether the change is clinically significant.

			Notes
			-----
			The function uses the following clinical significance thresholds:
			- Resistance: 15% change
			- Capacitance: 20% change
			- Absolute impedance (Z): 15% change

			If either value is zero or invalid, the comparison for that parameter is skipped.
		"""

		changes = {}
		clinically_significant = {}

		# Define significance thresholds based on clinical literature
		significance_thresholds = {
			'resistance': 0.15,  # 15% change is clinically significant
			'capacitance': 0.20, # 20% change is clinically significant
			'Z': 0.15  # 15% change is clinically significant for absolute impedance
		}

		freq_types = ['lowest_freq', 'center_freq', 'highest_freq']
		params = ['Z', 'resistance', 'capacitance']

		for freq_type in freq_types:
			current_freq_data = current_visit.get('sensor_data', {}).get('impedance', {}).get(freq_type, {})
			previous_freq_data = previous_visit.get('sensor_data', {}).get('impedance', {}).get(freq_type, {})

			for param in params:
				try:
					current_val = float(current_freq_data.get(param, 0))
					previous_val = float(previous_freq_data.get(param, 0))

					if previous_val != 0 and current_val != 0:
						percent_change = (current_val - previous_val) / previous_val
						key = f"{param}_{freq_type}"
						changes[key] = percent_change

						# Determine clinical significance
						is_significant = abs(percent_change) > significance_thresholds.get(param, 0.15)
						clinically_significant[key] = is_significant
				except (ValueError, TypeError, ZeroDivisionError):
					continue

		return changes, clinically_significant

	@staticmethod
	def calculate_tissue_health_index(visit: VisitsDataType) -> tuple[float | None, str]:
		"""
		Calculate a tissue health index based on bioimpedance measurements from a patient visit.

		This function analyzes impedance data at different frequencies to determine tissue health.
		It calculates a health score using the ratio of low frequency to high frequency impedance,
		and optionally incorporates phase angle data when available.

		Parameters:
		-----------
		visit : dict
			A dictionary containing visit data, with the following structure:
			{
				'sensor_data': {
					'impedance': {
						'lowest_freq': {
							'Z': float,  # Impedance magnitude at low frequency
							'frequency': float  # Actual frequency value
						},
						'highest_freq': {
							'Z': float,  # Impedance magnitude at high frequency
							'resistance': float,  # Optional
							'capacitance': float,  # Optional
							'frequency': float  # Optional, defaults to 80000 Hz if not provided
						}
					}
				}
			}

		Returns:
		--------
		tuple(float or None, str)
			A tuple containing:
			- health_score: A normalized score from 0-100 representing tissue health, or None if calculation fails
			- interpretation: A string describing the tissue health status or an error message

		Notes:
		------
		The health score is calculated based on:
		1. Low/high frequency impedance ratio (LF/HF ratio):
			- Optimal range is 5-12, with 8.5 being ideal
			- Scores decrease as the ratio deviates from this range

		2. Phase angle (when resistance and capacitance are available):
			- Calculated as arctan(1/(2πfRC))
			- Optimal range is 5-7 degrees
			- Higher values (up to 7 degrees) indicate better tissue health

		The final score is a weighted average: 60% from the ratio score and 40% from phase angle score.
		If phase angle cannot be calculated, only the ratio score is used.
		"""

		sensor_data    = visit.get('sensor_data', {})
		impedance_data = sensor_data.get('impedance', {})

		# Extract absolute impedance at different frequencies
		low_freq  = impedance_data.get('lowest_freq', {})
		high_freq = impedance_data.get('highest_freq', {})

		low_z  = low_freq.get('Z', 0)
		high_z = high_freq.get('Z', 0)

		if low_z is None or high_z is None:
			return None, "Insufficient data for tissue health calculation"

		low_z  = float(low_z)
		high_z = float(high_z)

		if low_z > 0 and high_z > 0:
			# Calculate low/high frequency ratio
			lf_hf_ratio = low_z / high_z

			# Calculate phase angle if resistance and reactance available
			phase_angle = None
			if 'resistance' in high_freq and 'capacitance' in high_freq:
				r = float(high_freq.get('resistance', 0))
				c = float(high_freq.get('capacitance', 0))
				if r > 0 and c > 0:
					# Approximate phase angle calculation
					# Using arctan(1/(2πfRC))
					f = float(high_freq.get('frequency', 80000))
					phase_angle = math.atan(1/(2 * math.pi * f * r * c)) * (180/math.pi)

			# Normalize scores to 0-100 scale
			# Typical healthy ratio range: 5-12
			# This formula is based on bioimpedance analysis principles in wound healing:
			# - LF/HF ratio typically ranges from 5-12 in healthy tissue
			# - Optimal ratio is around 8.5 (midpoint of healthy range)
			# - Scores decrease as ratio deviates from this range
			# Within 5-12 range: Linear scaling from 100 (at 5) to 0 (at 12)
			# Outside range: Penalty increases with distance from optimal ratio
			ratio_score = max(0, min(100, (1 - (lf_hf_ratio - 5) / 7) * 100)) if 5 <= lf_hf_ratio <= 12 else max(0, 50 - abs(lf_hf_ratio - 8.5) * 5)

			if phase_angle:
				# Typical healthy phase angle range: 5-7 degrees
				# Phase angle calculation logic:
				# - Typical healthy phase angle range: 5-7 degrees
				# - Linear scaling from 0 to 100 as phase angle increases from 0 to 7 degrees
				# - Values above 7 degrees are capped at 100
				# - This approach assumes that higher phase angles (up to 7 degrees) indicate better tissue health
				phase_score = max(0, min(100, (phase_angle / 7) * 100))
				health_score = (ratio_score * 0.6) + (phase_score * 0.4)  # Weighted average
			else:
				health_score = ratio_score

			# Interpretation
			if health_score >= 80:
				interpretation = "Excellent tissue health"
			elif health_score >= 60:
				interpretation = "Good tissue health"
			elif health_score >= 40:
				interpretation = "Moderate tissue health"
			elif health_score >= 20:
				interpretation = "Poor tissue health"
			else:
				interpretation = "Very poor tissue health"

			return health_score, interpretation

		# except (ValueError, TypeError, ZeroDivisionError):
		# 	pass
		return None, "Insufficient data for tissue health calculation"

	@staticmethod
	def analyze_healing_trajectory(visits: List[VisitsDataType]) -> dict:
		"""
			Analyzes the wound healing trajectory based on impedance data from multiple visits.

			This function performs a linear regression analysis on the high-frequency impedance values
			over time to determine if there's a significant trend indicating wound healing or deterioration.

			Parameters
			-----------
			visits : list
				List of visit dictionaries, each containing visit data including sensor readings.
				Each visit dictionary should have:
				- DataColumns.VISIT_DATE: date of the visit
				- 'sensor_data': dict containing 'impedance' data with 'highest_freq' values including a 'Z' value representing impedance measurement

			Returns
			--------
			dict
				A dictionary containing:
				- 'status': 'insufficient_data' if fewer than 3 valid measurements, 'analyzed' otherwise
				- 'slope': slope of the linear regression (trend direction)
				- 'r_squared': coefficient of determination (strength of linear relationship)
				- 'p_value': statistical significance of the slope
				- 'dates': list of dates with valid measurements
				- 'values': list of impedance values used in analysis
				- 'interpretation': Clinical interpretation of results as one of:
					- "Strong evidence of healing progression"
					- "Moderate evidence of healing progression"
					- "Potential deterioration detected"
					- "No significant trend detected"

			Notes:
			------
			Negative slopes indicate healing (decreasing impedance), while positive slopes
			may indicate deterioration. The function requires at least 3 valid impedance
			readings to perform analysis.
		"""

		if len(visits) < 3:
			return {"status": "insufficient_data"}

		VISIT_DATE_TAG = WoundDataProcessor.get_visit_date_tag(visits)

		dates, z_values = [], []
		for visit in visits:
			try:
				high_freq = visit.get('sensor_data', {}).get('impedance', {}).get('highest_freq', {})
				z_val = float(high_freq.get('Z', 0))
				if z_val > 0:
					z_values.append(z_val)
					dates.append(visit.get(VISIT_DATE_TAG))

			except (ValueError, TypeError):
				continue

		if len(z_values) < 3:
			return {"status": "insufficient_data"}

		# Convert dates to numerical values for regression
		x = np.arange(len(z_values))
		y = np.array(z_values)

		# Perform linear regression
		slope, intercept, r_value, p_value, std_err = stats.linregress(x, y)
		r_squared = r_value ** 2

		# Determine clinical significance of slope
		result = {
			"slope": slope,
			"r_squared": r_squared,
			"p_value": p_value,
			"dates": dates,
			"values": z_values,
			"status": "analyzed"
		}

		# Interpret the slope
		if slope < -0.5 and p_value < 0.05:
			result["interpretation"] = "Strong evidence of healing progression"
		elif slope < -0.2 and p_value < 0.10:
			result["interpretation"] = "Moderate evidence of healing progression"
		elif slope > 0.5 and p_value < 0.05:
			result["interpretation"] = "Potential deterioration detected"
		else:
			result["interpretation"] = "No significant trend detected"

		return result

	@staticmethod
	def analyze_frequency_response(visit: VisitsDataType) -> dict:
		"""
		Analyze the pattern of impedance across different frequencies to assess tissue composition.

		This function utilizes bioelectrical impedance analysis (BIA) principles to evaluate
		tissue characteristics based on the frequency-dependent response to electrical current.
		It focuses on two key dispersion regions:

		1. Alpha Dispersion (low to center frequency):
			- Occurs in the kHz range
			- Reflects extracellular fluid content and cell membrane permeability
			- Large alpha dispersion may indicate edema or inflammation
			- Formula: alpha_dispersion = (Z_low - Z_center) / Z_low

		2. Beta Dispersion (center to high frequency):
			- Occurs in the MHz range
			- Reflects cellular density and intracellular properties
			- Large beta dispersion may indicate increased cellular content or structural changes
			- Formula: beta_dispersion = (Z_center - Z_high) / Z_center

		The function calculates these dispersion ratios and interprets them to assess tissue
		composition, providing insights into potential edema, cellular density, or active
		tissue remodeling.

		The interpretation logic:
		- If alpha_dispersion > 0.4 and beta_dispersion < 0.2:
			"High extracellular fluid content, possible edema"
		- If alpha_dispersion < 0.2 and beta_dispersion > 0.3:
			"High cellular density, possible granulation"
		- If alpha_dispersion > 0.3 and beta_dispersion > 0.3:
			"Mixed tissue composition, active remodeling"
		- Otherwise:
			"Normal tissue composition pattern"

		Args:
			visit (dict): A dictionary containing visit data, including sensor readings.

		Returns:
		--------
		dict
			A dictionary containing dispersion characteristics and interpretation of tissue composition.
		"""
		results = {}
		sensor_data    = visit.get('sensor_data', {})
		impedance_data = sensor_data.get('impedance', {})

		# Extract absolute impedance at different frequencies
		low_freq    = impedance_data.get('lowest_freq', {})
		center_freq = impedance_data.get('center_freq', {})
		high_freq   = impedance_data.get('highest_freq', {})

		try:
			# Get impedance values
			z_low    = low_freq.get('Z')
			z_center = center_freq.get('Z')
			z_high   = high_freq.get('Z')

			# First check if all values are not None
			if z_low is not None and z_center is not None and z_high is not None:
				# Then convert to float
				z_low    = float(z_low)
				z_center = float(z_center)
				z_high   = float(z_high)

				if z_low > 0 and z_center > 0 and z_high > 0:
					# Calculate alpha dispersion (low-to-center frequency drop)
					alpha_dispersion = (z_low - z_center) / z_low

					# Calculate beta dispersion (center-to-high frequency drop)
					beta_dispersion = (z_center - z_high) / z_center

					results['alpha_dispersion'] = alpha_dispersion
					results['beta_dispersion']  = beta_dispersion

					# Interpret dispersion patterns
					if alpha_dispersion > 0.4 and beta_dispersion < 0.2:
						results['interpretation'] = "High extracellular fluid content, possible edema"
					elif alpha_dispersion < 0.2 and beta_dispersion > 0.3:
						results['interpretation'] = "High cellular density, possible granulation"
					elif alpha_dispersion > 0.3 and beta_dispersion > 0.3:
						results['interpretation'] = "Mixed tissue composition, active remodeling"
					else:
						results['interpretation'] = "Normal tissue composition pattern"
				else:
					results['interpretation'] = "Insufficient frequency data for analysis (zero values)"
			else:
				results['interpretation'] = "Insufficient frequency data for analysis (missing values)"
		except (ValueError, TypeError, ZeroDivisionError) as e:
			error_message = f"Error processing frequency response data: {type(e).__name__}: {str(e)}"
			print(error_message)  # For console debugging
			traceback.print_exc()  # Print the full traceback
			results['interpretation'] = error_message  # Or keep the generic message if preferred

		return results

	@staticmethod
	def detect_impedance_anomalies(previous_visits: List[VisitsDataType], current_visit: VisitsDataType, z_score_threshold: float = 2.5) -> dict:
		"""
		Detects anomalies in impedance measurements by comparing the current visit's
		values with historical data from previous visits.

		This function analyzes three key impedance parameters:
		1. High-frequency impedance (Z)
		2. Low-frequency resistance
		3. High-frequency capacitance

		It calculates z-scores for each parameter based on historical means and standard
		deviations. Values exceeding the specified z-score threshold trigger alerts
		with clinical interpretations of the observed changes.

		Parameters:
		-----------
		previous_visits : list
			List of dictionaries containing data from previous visits, each with a
			'sensor_data' field containing impedance measurements.
		current_visit : dict
			Dictionary containing the current visit data with a 'sensor_data' field
			containing impedance measurements.
		z_score_threshold : float, optional (default=2.5)
			Threshold for flagging anomalies. Values with absolute z-scores exceeding
			this threshold will be reported.

		Returns:
		--------
		dict
			A dictionary of alerts where keys are parameter identifiers and values are
			dictionaries containing:
			- 'parameter': Display name of the parameter
			- 'z_score': Calculated z-score for the current value
			- 'direction': Whether the anomaly is an 'increase' or 'decrease'
			- 'interpretation': Clinical interpretation of the observed change

		Notes:
		------
		- Requires at least 3 previous visits with valid measurements to establish baseline
		- Ignores parameters with missing or non-positive values
		- Provides different clinical interpretations based on parameter type and direction of change
		"""

		if len(previous_visits) < 3:
			return {}

		alerts = {}

		# Parameters to monitor
		params = [
					('highest_freq' , 'Z'           , 'High-frequency impedance')   ,
					('lowest_freq'  , 'resistance'  , 'Low-frequency resistance')   ,
					('highest_freq' , 'capacitance' , 'High-frequency capacitance')
				]

		current_impedance = current_visit.get('sensor_data', {}).get('impedance', {})

		for freq_type, param_name, display_name in params:
			# Collect historical values
			historical_values = []

			for visit in previous_visits:
				visit_impedance = visit.get('sensor_data', {}).get('impedance', {})
				freq_data = visit_impedance.get(freq_type, {})
				try:
					value = float(freq_data.get(param_name, 0))
					if value > 0:
						historical_values.append(value)
				except (ValueError, TypeError):
					continue

			if len(historical_values) >= 3:
				# Calculate historical statistics
				mean = np.mean(historical_values)
				std = np.std(historical_values)

				# Get current value
				current_freq_data = current_impedance.get(freq_type, {})
				try:
					current_value = float(current_freq_data.get(param_name, 0))
					if current_value > 0 and std > 0:
						z_score = (current_value - mean) / std

						if abs(z_score) > z_score_threshold:
							direction = "increase" if z_score > 0 else "decrease"

							# Clinical interpretation
							if freq_type == 'highest_freq' and param_name == 'Z':
								if direction == 'increase':
									clinical_meaning = "Possible deterioration in tissue quality or increased inflammation"
								else:
									clinical_meaning = "Possible improvement in cellular integrity or reduction in edema"
							elif freq_type == 'lowest_freq' and param_name == 'resistance':
								if direction == 'increase':
									clinical_meaning = "Possible decrease in extracellular fluid or improved barrier function"
								else:
									clinical_meaning = "Possible increase in extracellular fluid or breakdown of tissue barriers"
							elif freq_type == 'highest_freq' and param_name == 'capacitance':
								if direction == 'increase':
									clinical_meaning = "Possible increase in cellular density or membrane integrity"
								else:
									clinical_meaning = "Possible decrease in viable cell count or membrane dysfunction"
							else:
								clinical_meaning = "Significant change detected, clinical correlation advised"

							key = f"{freq_type}_{param_name}"
							alerts[key] = {
								"parameter": display_name,
								"z_score": z_score,
								"direction": direction,
								"interpretation": clinical_meaning
							}
				except (ValueError, TypeError):
					continue

		return alerts

	@staticmethod
	def assess_infection_risk(current_visit: VisitsDataType, previous_visit: VisitsDataType | None = None) -> dict:
		"""
			Evaluates the infection risk for a wound based on bioimpedance measurements.

			This function analyzes bioelectrical impedance data from the current visit and optionally
			compares it with a previous visit to determine infection risk. It considers multiple factors
			including impedance ratios, changes in resistance, and phase angle calculations.

			Parameters
			----------
			current_visit : dict
				Dictionary containing current visit data with sensor_data.impedance measurements
				(including lowest_freq and highest_freq values with Z, resistance, capacitance)
			previous_visit : dict, optional
				Dictionary containing previous visit data in the same format as current_visit,
				used for trend analysis

			Returns
			-------
			dict
				A dictionary containing:
				- risk_score: numeric score from 0-100 indicating infection risk
				- risk_level: string interpretation ("Low", "Moderate", or "High" infection risk)
				- contributing_factors: list of specific factors that contributed to the risk assessment

			Notes
			-----
			The function evaluates three primary factors:
			1. Low/high frequency impedance ratio (values >15 indicate increased risk)
			2. Sudden increase in low-frequency resistance
			3. Phase angle measurement

			The final risk score is capped between 0 and 100.
		"""


		risk_score = 0
		factors = []

		current_impedance = current_visit.get('sensor_data', {}).get('impedance', {})

		# Factor 1: Low/high frequency impedance ratio
		low_freq = current_impedance.get('lowest_freq', {})
		high_freq = current_impedance.get('highest_freq', {})

		try:
			low_z = float(low_freq.get('Z', 0))
			high_z = float(high_freq.get('Z', 0))

			if low_z > 0 and high_z > 0:
				ratio = low_z / high_z
				# Ratios > 15 are associated with increased infection risk in literature
				if ratio > 20:
					risk_score += 40
					factors.append("Very high low/high frequency impedance ratio")
				elif ratio > 15:
					risk_score += 25
					factors.append("Elevated low/high frequency impedance ratio")
		except (ValueError, TypeError, ZeroDivisionError):
			pass

		# Factor 2: Sudden increase in low-frequency resistance (inflammatory response)
		if previous_visit:
			prev_impedance = previous_visit.get('sensor_data', {}).get('impedance', {})
			prev_low_freq = prev_impedance.get('lowest_freq', {})

			try:
				current_r = float(low_freq.get('resistance', 0))
				prev_r = float(prev_low_freq.get('resistance', 0))

				if prev_r > 0 and current_r > 0:
					pct_change = (current_r - prev_r) / prev_r
					if pct_change > 0.30:  # >30% increase
						risk_score += 30
						factors.append("Significant increase in low-frequency resistance")
			except (ValueError, TypeError, ZeroDivisionError):
				pass

		# Factor 3: Phase angle calculation (if resistance and capacitance available)
		try:
			r = float(high_freq.get('resistance', 0))
			c = float(high_freq.get('capacitance', 0))
			f = float(high_freq.get('frequency', 80000))

			if r > 0 and c > 0:
				# Phase angle calculation based on the complex impedance model
				# It represents the phase difference between voltage and current in AC circuits
				# Lower phase angles indicate less healthy or more damaged tissue
				phase_angle = math.atan(1/(2 * math.pi * f * r * c)) * (180/math.pi)

				# Phase angle thresholds based on bioimpedance literature:
				# <2°: Indicates severe tissue damage or very poor health
				# 2-3°: Suggests compromised tissue health
				# >3°: Generally associated with healthier tissue
				if phase_angle < 2:
					risk_score += 30
					factors.append("Very low phase angle (<2°): Indicates severe tissue damage")
				elif phase_angle < 3:
					risk_score += 15
					factors.append("Low phase angle (2-3°): Suggests compromised tissue health")
		except (ValueError, TypeError, ZeroDivisionError):
			pass

		# Limit score to 0-100 range
		risk_score = min(100, max(0, risk_score))

		# Interpret risk level
		if risk_score >= 60:
			interpretation = "High infection risk"
		elif risk_score >= 30:
			interpretation = "Moderate infection risk"
		else:
			interpretation = "Low infection risk"

		return {
			"risk_score": risk_score,
			"risk_level": interpretation,
			"contributing_factors": factors
		}

	@staticmethod
	def calculate_cole_parameters(visit: VisitsDataType) -> dict:
		"""
		Calculate Cole-Cole parameters from impedance measurement data in a visit.

		This function extracts impedance data at low, center, and high frequencies from a
		visit dictionary and calculates key Cole-Cole model parameters including:
		- R0 (low frequency resistance)
		- Rinf (high frequency resistance)
		- Cm (membrane capacitance)
		- Tau (time constant)
		- Alpha (tissue heterogeneity index)
		- Tissue homogeneity interpretation

		Parameters
		----------
		visit : dict
			A dictionary containing visit data, with a nested 'sensor_data' dictionary
			that includes impedance measurements at different frequencies

		Returns
		-------
		dict
			A dictionary containing calculated Cole-Cole parameters and tissue homogeneity assessment.
			May be empty if required data is missing or calculations fail.

		Notes
		-----
		The function handles exceptions for missing data, type errors, value errors, and
		division by zero, returning whatever parameters were successfully calculated before
		the exception occurred.
		"""

		results = {}
		impedance_data = visit.get('sensor_data', {}).get('impedance', {})

		# Extract resistance at different frequencies
		low_freq = impedance_data.get('lowest_freq', {})
		center_freq = impedance_data.get('center_freq', {})
		high_freq = impedance_data.get('highest_freq', {})

		try:
			# Get resistance values
			r_low    = float(low_freq.get('resistance', 0))
			r_center = float(center_freq.get('resistance', 0))
			r_high   = float(high_freq.get('resistance', 0))

			# Get capacitance values
			c_center = float(center_freq.get('capacitance', 0))

			# Get frequency values
			f_center = float(center_freq.get('frequency', 7499))

			if r_low > 0 and r_high > 0:
				# Approximate R0 and R∞
				results['R0'] = r_low  # Low frequency resistance approximates R0
				results['Rinf'] = r_high  # High frequency resistance approximates R∞

				# Calculate membrane capacitance (Cm)
				if r_center > 0 and c_center > 0:
					results['Fc'] = f_center

					# Calculate time constant
					tau = 1 / (2 * math.pi * f_center)
					results['Tau'] = tau

					# Membrane capacitance estimation
					if (r_low - r_high) > 0 and r_high > 0:
						cm = tau / ((r_low - r_high) * r_high)
						results['Cm'] = cm

				# Calculate alpha (tissue heterogeneity)
				# Using resistance values to estimate alpha
				if r_low > 0 and r_center > 0 and r_high > 0:
					# Simplified alpha estimation
					alpha_est = 1 - (r_center / math.sqrt(r_low * r_high))
					results['Alpha'] = max(0, min(1, abs(alpha_est)))

					# Interpret alpha value
					if results['Alpha'] < 0.6:
						results['tissue_homogeneity'] = "High tissue homogeneity"
					elif results['Alpha'] < 0.8:
						results['tissue_homogeneity'] = "Moderate tissue homogeneity"
					else:
						results['tissue_homogeneity'] = "Low tissue homogeneity (heterogeneous tissue)"
		except (ValueError, TypeError, ZeroDivisionError, KeyError):
			pass

		return results

	@staticmethod
	def generate_clinical_insights(analysis):
		"""
			Generates clinical insights based on various wound analysis results.

			This function processes different aspects of wound analysis data and translates
			them into actionable clinical insights with corresponding confidence levels
			and recommendations.

			Parameters
			----------
			analyses : dict
				A dictionary containing analysis results with potential keys:
				- 'healing_trajectory': Contains slope, p_value, and status of wound healing over time
				- 'infection_risk': Contains risk_score and contributing_factors for infection
				- 'frequency_response': Contains interpretation of impedance frequency response
				- 'anomalies': Contains significant deviations in parameters with z-scores

			Returns
			-------
			list
				A list of dictionaries, each representing a clinical insight with keys:
				- 'insight': The main clinical observation
				- 'confidence': Level of certainty (High, Moderate, etc.)
				- 'recommendation': Suggested clinical action (when applicable)
				- 'supporting_factors' or 'clinical_meaning': Additional context (when available)

			Notes
			-----
			The function generates insights based on the following criteria:
			1. Healing trajectory based on impedance trends and statistical significance
			2. Infection risk assessment based on risk score
			3. Tissue composition analysis from frequency response data
			4. Anomaly detection for significant deviations in measured parameters
		"""

		insights = []

		# Healing trajectory insights
		if 'healing_trajectory' in analysis:
			trajectory = analysis['healing_trajectory']
			if trajectory.get('status') == 'analyzed':
				if trajectory.get('slope', 0) < -0.3 and trajectory.get('p_value', 1) < 0.05:
					insights.append({
						"insight": "Strong evidence of consistent wound healing progression based on impedance trends",
						"confidence": "High",
						"recommendation": "Continue current treatment protocol"
					})
				elif trajectory.get('slope', 0) > 0.3 and trajectory.get('p_value', 1) < 0.1:
					insights.append({
						"insight": "Potential stalling or deterioration in wound healing process",
						"confidence": "Moderate",
						"recommendation": "Consider reassessment of treatment approach"
					})

		# Infection risk insights
		if 'infection_risk' in analysis:
			risk = analysis['infection_risk']
			if risk.get('risk_score', 0) > 50:
				insights.append({
					"insight": f"Elevated infection risk detected ({risk.get('risk_score')}%)",
					"confidence": "Moderate to High",
					"recommendation": "Consider microbiological assessment and prophylactic measures",
					"supporting_factors": risk.get('contributing_factors', [])
				})

		# Tissue composition insights
		if 'frequency_response' in analysis:
			freq_response = analysis['frequency_response']
			if 'interpretation' in freq_response:
				insights.append({
					"insight": freq_response['interpretation'],
					"confidence": "Moderate",
					"recommendation": "Correlate with clinical assessment of wound bed"
				})

		# Anomaly detection insights
		if 'anomalies' in analysis and analysis['anomalies']:
			for param, anomaly in analysis['anomalies'].items():
				insights.append({
					"insight": f"Significant {anomaly.get('direction')} in {anomaly.get('parameter')} detected (z-score: {anomaly.get('z_score', 0):.2f})",
					"confidence": "High" if abs(anomaly.get('z_score', 0)) > 3 else "Moderate",
					"clinical_meaning": anomaly.get('interpretation', '')
				})

		return insights

	@staticmethod
	def classify_wound_healing_stage(analysis):
		"""
		Classifies the wound healing stage based on bioimpedance and tissue health analyses.

		The function determines whether the wound is in the Inflammatory, Proliferative,
		or Remodeling phase by analyzing tissue health scores, frequency response characteristics,
		and Cole parameters from bioimpedance measurements.

		Parameters
		----------
		analyses : dict
			A dictionary containing various wound analysis results with the following keys:
			- 'tissue_health': tuple (score, description) where score is a numerical value
			- 'frequency_response': dict with keys 'alpha_dispersion' and 'beta_dispersion'
			- 'cole_parameters': dict containing Cole-Cole model parameters

		Returns
		-------
		dict
			A dictionary with the following keys:
			- 'stage': str, the determined healing stage ('Inflammatory', 'Proliferative', or 'Remodeling')
			- 'characteristics': list, notable characteristics of the identified stage
			- 'confidence': str, confidence level of the classification ('Low', 'Moderate', or 'High')

		Notes
		-----
		The classification uses the following general criteria:
		- Inflammatory: High alpha dispersion, low tissue health score
		- Proliferative: High beta dispersion, moderate tissue health, moderate alpha dispersion
		- Remodeling: Low alpha dispersion, high tissue health score

		If insufficient data is available, the function defaults to the Inflammatory stage with Low confidence.
		"""

		# Default to inflammatory if we don't have enough data
		stage = "Inflammatory"
		characteristics = []
		confidence = "Low"

		# Get tissue health index
		tissue_health = analysis.get('tissue_health', (None, ""))
		health_score = tissue_health[0] if tissue_health else None

		# Get frequency response
		freq_response = analysis.get('frequency_response', {})
		alpha = freq_response.get('alpha_dispersion', 0)
		beta = freq_response.get('beta_dispersion', 0)

		# Get Cole parameters
		cole_params = analysis.get('cole_parameters', {})

		# Stage classification logic
		if health_score is not None and freq_response and cole_params:
			confidence = "Moderate"

			# Inflammatory phase characteristics:
			# - High alpha dispersion (high extracellular fluid)
			# - Low tissue health score
			# - High low/high frequency ratio
			if alpha > 0.4 and health_score < 40:
				stage = "Inflammatory"
				characteristics = [
					"High extracellular fluid content",
					"Low tissue health score",
					"Elevated cellular permeability"
				]
				confidence = "High" if alpha > 0.5 and health_score < 30 else "Moderate"

			# Proliferative phase characteristics:
			# - High beta dispersion (cellular proliferation)
			# - Moderate tissue health score
			# - Moderate alpha dispersion
			elif beta > 0.3 and 40 <= health_score <= 70 and 0.2 <= alpha <= 0.4:
				stage = "Proliferative"
				characteristics = [
					"Active cellular proliferation",
					"Increasing tissue organization",
					"Moderate extracellular fluid"
				]
				confidence = "High" if beta > 0.4 and health_score > 50 else "Moderate"

			# Remodeling phase characteristics:
			# - Low alpha dispersion (reduced extracellular fluid)
			# - High tissue health score
			# - Low variability in impedance
			elif alpha < 0.2 and health_score > 70:
				stage = "Remodeling"
				characteristics = [
					"Reduced extracellular fluid",
					"Improved tissue organization",
					"Enhanced barrier function"
				]
				confidence = "High" if alpha < 0.15 and health_score > 80 else "Moderate"

		return {
			"stage": stage,
			"characteristics": characteristics,
			"confidence": confidence
		}

	@staticmethod
	def prepare_population_stats(df: pd.DataFrame) -> tuple:
		"""
		Calculate average impedance statistics across different groupings from the wound data.

		This function groups wound data by visit number to calculate mean impedance values,
		and also calculates the average impedance by wound type.

		Parameters
		----------
		df : pd.DataFrame
			DataFrame containing wound data with columns for 'Visit Number', CN.WOUND_TYPE,
			and various impedance measurements (Z, Z', Z'').

		Returns
		-------
		tuple
			A tuple containing two DataFrames:
			- avg_impedance: DataFrame with average impedance components by visit number
			- avg_by_type: DataFrame with average impedance by wound type
		"""
		CN = DColumns(df=df)

		# Impedance components by visit
		avg_impedance = df.groupby(CN.VISIT_NUMBER, observed=False)[
			[CN.HIGHEST_FREQ_ABSOLUTE, CN.HIGHEST_FREQ_REAL, CN.HIGHEST_FREQ_IMAGINARY]
		].mean().reset_index()

		# Impedance by wound type
		avg_by_type = df.groupby(CN.WOUND_TYPE, observed=False)[CN.HIGHEST_FREQ_ABSOLUTE].mean().reset_index()

		return avg_impedance, avg_by_type

	@staticmethod
	def generate_clinical_analysis(current_visit: VisitsDataType, previous_visit: VisitsDataType | None = None) -> dict:
		"""
		Generate a comprehensive clinical analysis based on the current visit data and optionally compare with previous visit.

		This method analyzes wound data to produce metrics related to tissue health, infection risk,
		and frequency response. When previous visit data is provided, it also calculates changes
		between visits and identifies significant changes.

		Parameters
		----------
		current_visit : dict
			Data from the current wound assessment visit
		previous_visit : dict, optional
			Data from the previous wound assessment visit for comparison, default is None

		Returns
		-------
		dict
			A dictionary containing the following keys:
			- 'tissue_health': Index quantifying overall tissue health
			- 'infection_risk': Assessment of wound infection probability
			- 'frequency_response': Analysis of tissue composition
			- 'changes': Differences between current and previous visit (only if previous_visit provided)
			- 'significant_changes': Notable changes requiring attention (only if previous_visit provided)
		"""

		analysis = {}

		# Calculate tissue health index
		analysis['tissue_health'] = ImpedanceAnalyzer.calculate_tissue_health_index(visit=current_visit)

		# Assess infection risk
		analysis['infection_risk'] = ImpedanceAnalyzer.assess_infection_risk(current_visit=current_visit, previous_visit=previous_visit)

		# Analyze tissue composition (frequency response)
		analysis['frequency_response'] = ImpedanceAnalyzer.analyze_frequency_response(visit=current_visit)

		# Calculate changes since previous visit
		if previous_visit:
			analysis['changes'], analysis['significant_changes'] = ImpedanceAnalyzer.calculate_visit_changes( current_visit=current_visit, previous_visit=previous_visit )

		return analysis

	@staticmethod
	def generate_advanced_analysis(visits: List[VisitsDataType]) -> dict:
		"""
		This method performs comprehensive analysis of wound data across multiple patient visits,
		including healing trajectory assessment, anomaly detection, Cole parameter calculations,
		tissue health evaluation, infection risk assessment, frequency response analysis,
		clinical insight generation, and wound healing stage classification.

			visits (List[VisitsDataType]): List of visit dictionaries containing wound measurement data. At least 3 visits are required for full analysis.

			dict: Dictionary with advanced analysis results containing the following keys:
				- 'status': 'insufficient_data' if fewer than 3 visits are provided
				- 'healing_trajectory': Analysis of wound healing progression over time
				- 'anomalies': Detected impedance anomalies compared to previous visits
				- 'cole_parameters': Calculated Cole model parameters for the latest visit
				- 'tissue_health': Tissue health index from the most recent visit
				- 'infection_risk': Assessment of infection risk based on recent measurements
				- 'frequency_response': Analysis of bioimpedance frequency response
				- 'insights': Generated clinical insights based on all analyses
				- 'healing_stage': Classification of current wound healing stage

		Raises:
			None explicitly, but may propagate exceptions from called methods

		Note:
			This method requires at least 3 visits to generate a complete analysis.
			If fewer visits are provided, returns a dictionary with status 'insufficient_data'.
		"""
		if len(visits) < 3:
			return {'status': 'insufficient_data'}

		analysis = {}

		# Analyze healing trajectory
		analysis['healing_trajectory'] = ImpedanceAnalyzer.analyze_healing_trajectory(visits=visits)

		# Detect anomalies
		analysis['anomalies'] = ImpedanceAnalyzer.detect_impedance_anomalies( previous_visits=visits[:-1], current_visit=visits[-1] )

		# Calculate Cole parameters
		analysis['cole_parameters'] = ImpedanceAnalyzer.calculate_cole_parameters(visit=visits[-1])

		# Get tissue health from most recent visit
		analysis['tissue_health'] = ImpedanceAnalyzer.calculate_tissue_health_index(visit=visits[-1])

		# Get infection risk assessment
		analysis['infection_risk'] = ImpedanceAnalyzer.assess_infection_risk(
			current_visit=visits[-1], previous_visit=visits[-2] if len(visits) > 1 else None
		)

		# Get frequency response analysis
		analysis['frequency_response'] = ImpedanceAnalyzer.analyze_frequency_response(visit=visits[-1])

		# Generate clinical insights
		analysis['insights'] = ImpedanceAnalyzer.generate_clinical_insights(analysis=analysis)

		# Classify wound healing stage
		analysis['healing_stage'] = ImpedanceAnalyzer.classify_wound_healing_stage(analysis=analysis)

		return analysis


def load_env():
	# Load environment variables from .env file
	from dotenv import load_dotenv

	# Try to load environment variables from different possible locations
	env_paths = [
		pathlib.Path(__file__).parent.parent / '.env',  # Project root .env
		pathlib.Path.cwd() / '.env',                    # Current working directory
	]

	for env_path in env_paths:
		if env_path.exists():
			load_dotenv(dotenv_path=env_path)
			break
